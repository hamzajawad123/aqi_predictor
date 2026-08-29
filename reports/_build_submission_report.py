"""Build reports/final_report.docx. Run: python reports/_build_submission_report.py"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = Path(__file__).resolve().parent / "final_report.docx"
FIG = Path(__file__).resolve().parent / "_report_figures"
LOGO = ROOT / "10_pearls_logo.png"
BLACK = RGBColor(0x00, 0x00, 0x00)
FONT = "Times New Roman"
BOLD_MARK = re.compile(r"(\*\*.+?\*\*)")


def _rfonts(run, name=FONT):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), name)


def set_run(run, *, size=12, bold=False, italic=False):
    _rfonts(run)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = BLACK


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(12)
    normal.font.color.rgb = BLACK
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    pf = normal.paragraph_format
    pf.space_after = Pt(8)
    pf.line_spacing = 1.15
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    heading_sizes = {1: 14, 2: 13, 3: 12}
    for level, size in heading_sizes.items():
        style = doc.styles[f"Heading {level}"]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLACK
        style.font.italic = False
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rFonts.set(qn(attr), FONT)
        color = rPr.find(qn("w:color"))
        if color is None:
            color = OxmlElement("w:color")
            rPr.append(color)
        color.set(qn("w:val"), "000000")
        style.paragraph_format.space_before = Pt(16 if level == 1 else 12)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.keep_with_next = True

    try:
        cap = doc.styles["Caption"]
        cap.font.name = FONT
        cap.font.size = Pt(11)
        cap.font.italic = True
        cap.font.color.rgb = BLACK
        cap.font.bold = False
        rPr = cap.element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rFonts.set(qn(attr), FONT)
    except KeyError:
        pass


def add_page_number(paragraph):
    run = paragraph.add_run()
    set_run(run, size=11)
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._element.append(fld1)
    run._element.append(instr)
    run._element.append(fld2)


def add_mixed_runs(paragraph, text, *, size=12, italic=False):
    """Turn **this** into bold."""
    for part in BOLD_MARK.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            set_run(run, size=size, bold=True, italic=italic)
        else:
            run = paragraph.add_run(part)
            set_run(run, size=size, italic=italic)


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run(run, size={1: 14, 2: 13, 3: 12}[level], bold=True)
        run.font.color.rgb = BLACK
    return h


def front_heading(doc, text):
    """Title used in the contents pages. Not a Heading, so it stays off the TOC."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(text)
    set_run(run, size=16, bold=True)
    return p


def toc_line(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.first_line_indent = Cm(0)
    if level == 1:
        p.paragraph_format.left_indent = Cm(0)
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run(text)
        set_run(run, size=12, bold=True)
    elif level == 2:
        p.paragraph_format.left_indent = Cm(0.85)
        run = p.add_run(text)
        set_run(run, size=12)
    else:
        p.paragraph_format.left_indent = Cm(1.6)
        run = p.add_run(text)
        set_run(run, size=11)
    return p


def body(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.first_line_indent = Cm(0)
    add_mixed_runs(p, text, size=12)
    return p


def centered(doc, text, *, size=12, bold=False, italic=False, space_after=8, space_before=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, italic=italic)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        for run in list(p.runs):
            run.text = ""
        add_mixed_runs(p, item, size=12)


def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        for run in list(p.runs):
            run.text = ""
        add_mixed_runs(p, item, size=12)


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.first_line_indent = Cm(0)
    try:
        p.style = doc.styles["Caption"]
    except KeyError:
        pass
    if ": " in text:
        label, rest = text.split(": ", 1)
        run = p.add_run(label + ": ")
        set_run(run, size=11, bold=True, italic=True)
        add_mixed_runs(p, rest, size=11, italic=True)
    else:
        add_mixed_runs(p, text, size=11, italic=True)
    return p


def add_picture(doc, name, width_in=6.3):
    path = FIG / name
    if not path.exists():
        body(doc, f"[Figure missing: {name}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:color"), "000000")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def shade_cell(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run(run, size=12, bold=True)
        shade_cell(hdr[i], "E6E6E6")
        set_cell_border(hdr[i])
    for r_i, row in enumerate(rows):
        cells = table.rows[r_i + 1].cells
        for c_i, val in enumerate(row):
            cells[c_i].text = ""
            p = cells[c_i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_run(run, size=12)
            set_cell_border(cells[c_i])
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_logo(doc, width_in=2.7):
    if not LOGO.exists():
        body(doc, "[10Pearls logo not found: 10_pearls_logo.png]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run()
    run.add_picture(str(LOGO), width=Inches(width_in))


def setup_sections(doc: Document):
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.different_first_page_header_footer = True
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run("Pearls AQI Predictor  |  Page ")
        set_run(run, size=11)
        add_page_number(fp)
        first = section.first_page_footer
        first.is_linked_to_previous = False
        if first.paragraphs:
            first.paragraphs[0].text = ""


# First training notebook scores
AT8_24 = [
    ("Persistence Baseline", "66.730405", "36.111431", "0.745204"),
    ("XGBoost", "74.364967", "50.019281", "0.683567"),
    ("LightGBM", "76.456676", "52.289149", "0.665515"),
    ("Ensemble_top3_tabular", "77.085201", "53.331851", "0.659993"),
    ("RandomForest", "84.914023", "59.164036", "0.587424"),
    ("Ridge", "87.584091", "63.751444", "0.561069"),
    ("GRU", "219.023846", "146.531186", "-1.740176"),
    ("LSTM", "232.032646", "156.863391", "-2.075345"),
    ("Prophet", "2167.324472", "1736.633538", "-267.777676"),
]
AT8_48 = [
    ("Persistence Baseline", "84.796609", "45.775352", "0.588437"),
    ("XGBoost", "97.213262", "69.199877", "0.459083"),
    ("Ensemble_top3_tabular", "98.112376", "70.283343", "0.449031"),
    ("LightGBM", "101.157242", "72.808352", "0.414302"),
    ("Ridge", "101.382767", "73.917026", "0.411688"),
    ("RandomForest", "118.050540", "86.744712", "0.202344"),
    ("LSTM", "234.142256", "152.548810", "-2.132563"),
    ("GRU", "246.401500", "157.766790", "-2.469181"),
    ("Prophet", "2159.443832", "1727.409530", "-265.909192"),
]
AT8_72 = [
    ("Persistence Baseline", "92.456124", "51.518059", "0.510480"),
    ("Ridge", "104.353359", "77.654450", "0.376391"),
    ("Ensemble_top3_tabular", "107.619632", "79.190920", "0.336742"),
    ("XGBoost", "111.990073", "82.949948", "0.281779"),
    ("LightGBM", "112.211034", "81.688058", "0.278942"),
    ("RandomForest", "116.121902", "85.821522", "0.227804"),
    ("LSTM", "235.968967", "160.365461", "-2.182948"),
    ("GRU", "239.031295", "159.196560", "-2.266098"),
    ("Prophet", "2150.783604", "1715.485984", "-263.906508"),
]


def metric_rows(pairs):
    return [[m, rmse, mae, r2] for m, rmse, mae, r2 in pairs]


def build():
    doc = Document()
    configure_styles(doc)
    setup_sections(doc)

    # Title page
    for _ in range(2):
        doc.add_paragraph()
    add_logo(doc, 2.7)
    centered(doc, "10Pearls Data Science Internship", size=14, bold=True, space_after=8)
    centered(doc, "Project Report", size=14, bold=True, space_after=20)
    centered(
        doc,
        "Pearls AQI Predictor:",
        size=16,
        bold=True,
        space_after=4,
    )
    centered(
        doc,
        "A Feature–Training–Inference Pipeline for",
        size=16,
        bold=True,
        space_after=4,
    )
    centered(
        doc,
        "Three-Day-Ahead Air Quality Forecasting in Lahore",
        size=16,
        bold=True,
        space_after=14,
    )
    centered(
        doc,
        "Continuous US EPA AQI computed from PM2.5  |  Horizons 24 / 48 / 72 hours",
        size=12,
        italic=True,
        space_after=22,
    )
    add_table(
        doc,
        ["Item", "Details"],
        [
            ["Project", "Pearls AQI Predictor"],
            ["Programme", "10Pearls Data Science Internship"],
            ["Author", "Hamza Jawad"],
            ["Target city", "Lahore (31.5497, 74.3436)"],
            ["GitHub", "https://github.com/hamzajawad123/aqi_predictor"],
            ["Live dashboard", "https://lahore-aqi-predictor.streamlit.app"],
            ["Document date", "29 August 2026"],
        ],
    )
    caption(doc, "Table 1: Project information.")
    doc.add_page_break()

    # Contents
    front_heading(doc, "Table of Contents")
    for title, level in [
        ("1. Executive Summary", 1),
        ("2. Introduction", 1),
        ("3. Problem Statement", 1),
        ("4. Project Objectives", 1),
        ("5. Project Scope", 1),
        ("6. Project Workflow", 1),
        ("7. Dataset Description", 1),
        ("7.1 Raw merged snapshot", 2),
        ("7.2 Data sources", 2),
        ("7.3 Target variable", 2),
        ("7.4 Feature-group table used in training", 2),
        ("8. Data Understanding", 1),
        ("9. Data Preprocessing", 1),
        ("9.1 Validation of raw merged data", 2),
        ("9.2 Hourly grid and short-gap interpolation", 2),
        ("9.3 Feature transforms", 2),
        ("9.4 Scaling", 2),
        ("9.5 Train / validation / test split", 2),
        ("9.6 What I did not do in the first training run", 2),
        ("10. Exploratory Data Analysis", 1),
        ("10.1 AQI distribution", 2),
        ("10.2 Pollutant distributions", 2),
        ("10.3 AQI over time", 2),
        ("10.4 AQI versus weather", 2),
        ("10.5 Correlation heatmap", 2),
        ("10.6 Autocorrelation", 2),
        ("10.7 Seasonal decomposition", 2),
        ("10.8 Smog season versus normal season", 2),
        ("11. Feature and Variable Analysis", 1),
        ("12. Methodology", 1),
        ("13. Approaches Considered and Used", 1),
        ("13.1 Persistence baseline", 2),
        ("13.2 Optuna-tuned tabular regression on deltas", 2),
        ("13.3 Mean ensemble of top-3 tabular models", 2),
        ("13.4 Sequence models (LSTM and GRU)", 2),
        ("13.5 Prophet on absolute AQI", 2),
        ("13.6 Production training window", 2),
        ("13.7 Classification of OpenWeather 1–5", 2),
        ("14. Models Used", 1),
        ("14.1 Persistence Baseline", 2),
        ("14.2 Ridge", 2),
        ("14.3 RandomForest", 2),
        ("14.4 XGBoost", 2),
        ("14.5 LightGBM", 2),
        ("14.6 Ensemble_top3_tabular", 2),
        ("14.7 LSTM", 2),
        ("14.8 GRU", 2),
        ("14.9 Prophet", 2),
        ("15. Model Training", 1),
        ("16. Model Evaluation", 1),
        ("17. Results", 1),
        ("17.1 Complete comparison (24 hours)", 2),
        ("17.2 Complete comparison (48 hours)", 2),
        ("17.3 Complete comparison (72 hours)", 2),
        ("18. Model Comparison", 1),
        ("19. Best Model on the First Training Notebook", 1),
        ("20. Production Training Window", 1),
        ("20.1 SHAP analysis of registered models", 2),
        ("20.1.1 24-hour registered model (XGBoost)", 3),
        ("20.1.2 48-hour registered model (RandomForest)", 3),
        ("20.1.3 72-hour registered ensemble members", 3),
        ("21. Classification Experiment", 1),
        ("21.1 Full-history classification", 2),
        ("21.2 Post-break classification", 2),
        ("22. Final Approach Selection", 1),
        ("23. Why Other Models Were Not Selected", 1),
        ("23.1 Ridge", 2),
        ("23.2 RandomForest", 2),
        ("23.3 XGBoost", 2),
        ("23.4 LightGBM", 2),
        ("23.5 Ensemble_top3_tabular", 2),
        ("23.6 LSTM", 2),
        ("23.7 GRU", 2),
        ("23.8 Prophet", 2),
        ("24. Serving, Dashboard and Automation", 1),
        ("25. Key Findings", 1),
        ("26. Limitations", 1),
        ("27. Conclusion", 1),
    ]:
        toc_line(doc, title, level)

    front_heading(doc, "List of Figures")
    for line in [
        "Figure 1: Distribution of AQI (histogram with KDE) and boxplot",
        "Figure 2: Histograms of the eight pollutant series",
        "Figure 3: AQI over time for Lahore",
        "Figure 4: AQI versus temperature, humidity, wind speed and pressure",
        "Figure 5: Correlation heatmap of numeric raw columns",
        "Figure 6: ACF and PACF of hourly AQI",
        "Figure 7: Seasonal decomposition of hourly AQI (period = 24)",
        "Figure 8: AQI in smog season versus the rest of the year",
        "Figure 9: Test RMSE by model at 24 hours",
        "Figure 10: Test RMSE by model at 48 hours",
        "Figure 11: Test RMSE by model at 72 hours",
        "Figure 12: SHAP beeswarm, 24h XGBoost",
        "Figure 13: Mean |SHAP| bars, 24h XGBoost",
        "Figure 14: SHAP beeswarm, 48h RandomForest",
        "Figure 15: Mean |SHAP| bars, 48h RandomForest",
        "Figure 16: SHAP beeswarm, 72h XGBoost_L1",
        "Figure 17: Mean |SHAP| bars, 72h XGBoost_L1",
        "Figure 18: SHAP beeswarm, 72h RandomForest",
        "Figure 19: Mean |SHAP| bars, 72h RandomForest",
        "Figure 20: SHAP beeswarm, 72h LightGBM",
        "Figure 21: Mean |SHAP| bars, 72h LightGBM",
    ]:
        toc_line(doc, line, 2)

    front_heading(doc, "List of Tables")
    for line in [
        "Table 1: Project information",
        "Table 2: Raw snapshot used for EDA",
        "Table 3: Raw snapshot columns and dtypes",
        "Table 4: PM2.5 breakpoints used to compute AQI",
        "Table 5: Feature group and split from the first training notebook",
        "Table 6: AQI descriptive statistics and ADF test",
        "Table 7: Yearly mean AQI",
        "Table 8: AQI grouped by smog-season flag",
        "Table 9: Correlation with AQI",
        "Table 10: Validation ranges",
        "Table 11: Persistence baseline test metrics",
        "Table 12: Ridge test metrics",
        "Table 13: RandomForest test metrics",
        "Table 14: XGBoost test metrics",
        "Table 15: LightGBM test metrics",
        "Table 16: Ensemble_top3_tabular test metrics",
        "Table 17: LSTM test metrics",
        "Table 18: GRU test metrics",
        "Table 19: Prophet test metrics",
        "Table 20: Registry summary from the first training notebook",
        "Table 21: Full test comparison at 24 hours",
        "Table 22: Full test comparison at 48 hours",
        "Table 23: Full test comparison at 72 hours",
        "Table 24: Persistence versus lowest-RMSE learned model",
        "Table 25: Production-window selected models",
        "Table 26: Production-window persistence versus selected RMSE",
        "Table 27: Top 10 mean |SHAP| features, 24h XGBoost",
        "Table 28: Top 10 mean |SHAP| features, 48h RandomForest",
        "Table 29: Highest mean |SHAP| feature per 72h ensemble member",
        "Table 30: Classification on full history (1–5 labels)",
        "Table 31: Classification on the post-4 April 2025 window (1–5 labels)",
        "Table 32: Classification accuracy and macro F1 (summary)",
        "Table 33: Why I selected continuous regression on the post-break window",
    ]:
        toc_line(doc, line, 2)

    doc.add_page_break()

    # Executive summary
    heading(doc, "1. Executive Summary", 1)
    body(
        doc,
        "I built a three-day-ahead Air Quality Index forecast for Lahore as my 10Pearls Data "
        "Science Internship project. Pollution concentrations come from the OpenWeather Air "
        "Pollution API. Weather comes from Open-Meteo. I store engineered features and trained "
        "models in Hopsworks, optionally wrap inference in FastAPI, and show the forecast on a "
        "Streamlit dashboard at https://lahore-aqi-predictor.streamlit.app."
    )
    body(
        doc,
        "The target I predict is a **continuous US EPA AQI from PM2.5**, not OpenWeather’s "
        "1–5 category field. Most models learn the change in AQI over 24, 48 or 72 hours and "
        "I add that change back to the current AQI before scoring. I used RMSE, MAE and R² on "
        "the reconstructed absolute AQI. A model only goes into the Hopsworks registry if it "
        "beats a persistence baseline (tomorrow’s AQI equals today’s) on **all three** metrics."
    )
    body(
        doc,
        "My first full training notebook read feature group **aqi_features v3** (46,100 rows, "
        "42 input features, December 2020 to August 2026). I used a season-aligned split: "
        "train to 1 June 2024 (27,947 rows), validation to 1 June 2025 (8,075 rows), and test "
        "from then on (10,078 rows). Optuna had 15 trials per tabular model. "
        "**Persistence was best at every horizon.** No learned model beat it on RMSE, MAE and "
        "R² together, so nothing was registered from that run."
    )
    body(
        doc,
        "I later restricted training to **4 April 2025** onwards after noticing a break in "
        "OpenWeather’s AQI series (mean hourly |AQI change| drops from about 46 to about 4.5). "
        "On that shorter window the production models did beat persistence (RMSE **28.88 / 30.54 / "
        "31.65** at 24 / 48 / 72 hours). I also tried classifying OpenWeather’s 1–5 category instead "
        "of continuous AQI; that path did not beat “today’s class” in a stable way, so I did not "
        "adopt it. **The production choice is continuous EPA AQI regression on the post-4 April 2025 "
        "window.** SHAP for those registered models is in Section 20.1. Section 17 is the first "
        "notebook; Section 20 is the later regression window; Section 21 is classification; "
        "Section 22 is the final choice."
    )

    heading(doc, "2. Introduction", 1)
    body(
        doc,
        "The project is a serverless Feature / Training / Inference pipeline. The flow I "
        "implemented is:"
    )
    bullets(
        doc,
        [
            "OpenWeather (pollution) and Open-Meteo (weather) feed the feature pipeline.",
            "The feature pipeline writes to the Hopsworks Feature Store.",
            "The training pipeline writes to the Hopsworks Model Registry.",
            "FastAPI can serve a prediction; it is optional.",
            "Streamlit reads Hopsworks directly and shows the forecast.",
        ],
    )
    body(
        doc,
        "Default location is Lahore at **31.5497, 74.3436**. History starts on **27 November 2020**, "
        "the first day OpenWeather’s free air-pollution archive is available. I forecast 24, 48 "
        "and 72 hours. Smog season in this project is October through January."
    )
    body(
        doc,
        "On **4 April 2025** the OpenWeather series changes character: mean hourly |AQI change| "
        "collapses from about 46 to about 4.5 and does not recover, with the same break in PM2.5. "
        "I now default training to that date. The first notebook did not apply that filter; it "
        "trained on the full v3 table described in Section 9."
    )

    heading(doc, "3. Problem Statement", 1)
    body(
        doc,
        "The brief is to forecast a **continuous AQI** for Lahore at 24, 48 and 72 hours. "
        "OpenWeather’s native `main.aqi` field is a coarse 1–5 category. RMSE, MAE and R² need "
        "a numeric target with meaningful distances, so I compute US EPA AQI from PM2.5 instead "
        "and keep the 1–5 field only as a reference column."
    )
    body(
        doc,
        "Hourly AQI is strongly autocorrelated, so “AQI in N hours equals AQI now” is a serious "
        "baseline, not a toy. I only register a model if it is strictly better on RMSE, strictly "
        "better on MAE, and strictly better on R² at the same time."
    )

    heading(doc, "4. Project Objectives", 1)
    bullets(
        doc,
        [
            "Forecast Lahore AQI at **24, 48 and 72 hours**.",
            "Compare continuous regression (0–500 EPA AQI) with classification of OpenWeather’s 1–5 category, and keep only the approach that beats a serious baseline.",
            "Use one pollution source (OpenWeather) and one weather source (Open-Meteo) for both history and the hourly path.",
            "Store engineered features in Hopsworks (`aqi_features`).",
            "Register `aqi_forecaster_{24,48,72}h` only when a model beats persistence on RMSE, MAE and R².",
            "Serve forecasts through FastAPI when needed, and present them in Streamlit.",
            "Automate hourly feature updates and daily retraining with GitHub Actions.",
        ],
    )

    heading(doc, "5. Project Scope", 1)
    bullets(
        doc,
        [
            "City: Lahore at the default coordinates. The pipeline is location-parameterised; I did not run a multi-city training experiment.",
            "AQI is computed from **PM2.5 only**, using the 2024 US EPA PM2.5 breakpoint table. A full six-pollutant EPA maximum is out of scope.",
            "I fetch OpenWeather’s 4-day pollution forecast for possible later comparison, but I do not use it in training.",
            "The first training notebook covers Ridge, RandomForest, XGBoost, LightGBM, a top-3 tabular mean ensemble, LSTM, GRU, Prophet, and persistence.",
            "`/predict` serves tabular or ensemble payloads that consume one feature row. Prophet needs a date series; LSTM/GRU need a 24-hour window, so those are not served that way.",
            "Production feature-group version is **4**. The first notebook used version **3**.",
        ],
    )

    heading(doc, "6. Project Workflow", 1)
    numbered(
        doc,
        [
            "Define the problem: continuous EPA AQI from PM2.5 at 24/48/72 hours for Lahore.",
            "Collect OpenWeather pollution from 27 November 2020 and Open-Meteo weather for the same hours; merge on UTC timestamp.",
            "Write a local raw snapshot (`aqi_raw_merged.parquet`).",
            "Validate required columns, duplicates, nulls and physical bounds before insert.",
            "Explore the raw snapshot (not the engineered feature group).",
            "Engineer features on a strict hourly grid: time encodings, log1p pollutants, lags, rolling windows, change rates, weather interactions, absolute and delta targets.",
            "Write the feature group to Hopsworks (v3 in the first notebook; v4 in current config).",
            "Train and evaluate RMSE, MAE and R² on absolute future AQI versus persistence.",
            "Register only if the model beats persistence on all three metrics.",
            "Serve through Streamlit (and optional FastAPI) reading Hopsworks.",
        ],
    )

    heading(doc, "7. Dataset Description", 1)
    heading(doc, "7.1 Raw merged snapshot", 2)
    body(
        doc,
        "I explored the local raw merge of pollution and weather. After loading I added month "
        "and a smog-season flag for plots; those two columns are not stored in the parquet."
    )
    add_table(
        doc,
        ["Item", "Details from local raw snapshot"],
        [
            ["File", "aqi_raw_merged.parquet"],
            ["Rows", "46419"],
            ["Stored columns", "16"],
            ["Timestamp range", "2020-11-27 00:00:00 to 2026-08-14 13:00:00"],
            ["Duplicate timestamps", "0"],
            ["Missing values (all stored columns)", "none"],
            ["Target used in this project", "aqi (US EPA AQI from PM2.5)"],
            ["Reference column (not the target)", "openweather_aqi_category"],
        ],
    )
    caption(doc, "Table 2: Raw snapshot used for EDA figures and descriptive statistics.")
    body(
        doc,
        "Stored columns and data types, before I added the two EDA-only columns:"
    )
    add_table(
        doc,
        ["Column", "dtype on snapshot"],
        [
            ["timestamp", "datetime64[ns]"],
            ["openweather_aqi_category", "float64"],
            ["co", "float64"],
            ["no", "float64"],
            ["no2", "float64"],
            ["o3", "float64"],
            ["so2", "float64"],
            ["pm2_5", "float64"],
            ["pm10", "float64"],
            ["nh3", "float64"],
            ["aqi", "int64"],
            ["temperature", "float64"],
            ["humidity", "int64"],
            ["wind_speed", "float64"],
            ["wind_deg", "int64"],
            ["pressure", "float64"],
        ],
    )
    caption(doc, "Table 3: Raw snapshot columns and dtypes.")
    heading(doc, "7.2 Data sources", 2)
    body(
        doc,
        "Pollution and AQI come from the OpenWeather Air Pollution API (current, forecast and "
        "historical). Weather comes from Open-Meteo’s archive and forecast APIs. I set the start "
        "date to 27 November 2020 so both sources share the first day. The Open-Meteo hourly "
        "fields I keep are temperature, humidity, wind speed, wind direction and surface pressure. "
        "All timestamps are UTC before the merge."
    )
    heading(doc, "7.3 Target variable", 2)
    body(
        doc,
        "I compute `aqi` from PM2.5 using the 2024-revised US EPA PM2.5 breakpoint table:"
    )
    add_table(
        doc,
        ["PM2.5 low (µg/m³)", "PM2.5 high (µg/m³)", "AQI low", "AQI high", "Label in source comment"],
        [
            ["0.0", "9.0", "0", "50", "Good"],
            ["9.1", "35.4", "51", "100", "Moderate"],
            ["35.5", "55.4", "101", "150", "Unhealthy for Sensitive Groups"],
            ["55.5", "125.4", "151", "200", "Unhealthy"],
            ["125.5", "225.4", "201", "300", "Very Unhealthy"],
            ["225.5", "325.4", "301", "500", "Hazardous"],
        ],
    )
    caption(doc, "Table 4: PM2.5 breakpoints used to compute AQI.")
    body(
        doc,
        "Concentration is truncated (not rounded) to one decimal place before lookup, matching "
        "EPA methodology. Values above 325.4 µg/m³ are extrapolated from the last band rather "
        "than hard-capped at 500. Official EPA AQI is the maximum across six criteria pollutants; "
        "I use **PM2.5 only**."
    )
    heading(doc, "7.4 Feature-group table used in training", 2)
    body(
        doc,
        "The first training notebook read the following from Hopsworks:"
    )
    add_table(
        doc,
        ["Item", "Value"],
        [
            ["Feature group", "aqi_features v3"],
            ["Shape", "(46100, 52)"],
            ["Range", "2020-12-05 18:00:00 -> 2026-08-08 06:00:00"],
            ["Delta / target columns confirmed", "aqi_delta_24h, aqi_delta_48h, aqi_delta_72h, aqi_target_24h, aqi_target_48h, aqi_target_72h"],
            ["n_features after dropping timestamp, hour, month, openweather_aqi_category and all six target columns", "42"],
            ["Train rows", "27947"],
            ["Validation rows", "8075"],
            ["Test rows", "10078"],
            ["Train window printed", "2020-12-05->2024-06-01"],
            ["Validation window printed", "2024-06-01->2025-06-01"],
            ["Test window printed", "2025-06-01->2026-08-08"],
        ],
    )
    caption(doc, "Table 5: Feature group and split from the first training notebook.")
    body(
        doc,
        "Version 3 has delta targets but was built on a gappy frame, so about **13%** of lags "
        "and targets spanned the wrong number of hours. Version 4 is the same features on a "
        "strict hourly grid. The first notebook used v3; production now defaults to v4."
    )

    heading(doc, "8. Data Understanding", 1)
    body(
        doc,
        "On the raw snapshot I used for EDA:"
    )
    add_table(
        doc,
        ["Statistic", "Value"],
        [
            ["aqi mean", "247.7509855877981"],
            ["aqi median", "177.0"],
            ["aqi standard deviation", "194.00734285333604"],
            ["aqi skewness", "1.907488520950293"],
            ["aqi kurtosis", "3.1202724556139034"],
            ["IQR outlier threshold (Q3 + 1.5 IQR)", "479.0"],
            ["Count of aqi values above that threshold", "5837"],
            ["ADF statistic", "-11.764392029744617"],
            ["ADF p-value", "1.1280420946421232e-21"],
        ],
    )
    caption(doc, "Table 6: AQI descriptive statistics and ADF test on the raw snapshot.")
    body(
        doc,
        "Yearly mean AQI after resampling to year-end:"
    )
    add_table(
        doc,
        ["Year-end timestamp", "Mean aqi"],
        [
            ["2020-12-31", "438.785818"],
            ["2021-12-31", "263.904701"],
            ["2022-12-31", "277.071900"],
            ["2023-12-31", "280.052722"],
            ["2024-12-31", "251.599010"],
            ["2025-12-31", "202.854102"],
            ["2026-12-31", "166.540349"],
        ],
    )
    caption(
        doc,
        "Table 7: Yearly mean AQI on the raw snapshot. 2020 starts on 27 November; 2026 ends on 14 August.",
    )
    body(
        doc,
        "I treat October–January as smog season. On the raw snapshot:"
    )
    add_table(
        doc,
        ["is_smog_season", "count", "mean", "std", "min", "25%", "50%", "75%", "max"],
        [
            ["0 (Normal)", "32406.0", "182.315405", "126.145572", "0.0", "112.0", "159.0", "193.0", "1000.0"],
            ["1 (Smog Oct–Jan)", "14013.0", "399.075145", "234.698541", "52.0", "207.0", "310.0", "553.0", "1000.0"],
        ],
    )
    caption(doc, "Table 8: AQI grouped by smog-season flag on the raw snapshot.")
    body(
        doc,
        "Pearson correlation of numeric columns with AQI (month excluded from the heatmap):"
    )
    add_table(
        doc,
        ["Variable", "Correlation with aqi"],
        [
            ["aqi", "1.000000"],
            ["pm2_5", "0.986036"],
            ["pm10", "0.958184"],
            ["co", "0.793539"],
            ["no", "0.560187"],
            ["no2", "0.546778"],
            ["is_smog_season", "0.512918"],
            ["pressure", "0.511338"],
            ["openweather_aqi_category", "0.477453"],
            ["so2", "0.367937"],
            ["nh3", "0.361643"],
            ["humidity", "0.284877"],
            ["wind_deg", "0.092419"],
            ["wind_speed", "-0.240057"],
            ["o3", "-0.266326"],
            ["temperature", "-0.545787"],
        ],
    )
    caption(doc, "Table 9: Correlation with aqi on the raw snapshot.")
    body(
        doc,
        "Pollutant skewness, sorted descending: no 3.92, nh3 3.01, so2 2.46, no2 2.27, co 1.99, "
        "pm2_5 1.41, pm10 1.29, o3 1.26. That right-skew is why I apply log1p to those eight "
        "pollutant columns before training."
    )

    heading(doc, "9. Data Preprocessing", 1)
    heading(doc, "9.1 Validation of raw merged data", 2)
    body(
        doc,
        "I validate the merge before writing to Hopsworks. Required columns are timestamp, aqi, "
        "pm2_5, pm10, co, no2, o3, so2, nh3, temperature, humidity, wind_speed, wind_deg and "
        "pressure. Duplicate timestamps are dropped (keep first). Rows with nulls in required "
        "columns are dropped. Out-of-range rows are dropped using these bounds:"
    )
    add_table(
        doc,
        ["Column", "Accepted range in VALID_RANGES"],
        [
            ["aqi", "0 to 1000"],
            ["pm2_5", "0 to 2000 µg/m³"],
            ["pm10", "0 to 2000 µg/m³"],
            ["co", "0 to 50000 µg/m³"],
            ["no2", "0 to 2000 µg/m³"],
            ["o3", "0 to 2000 µg/m³"],
            ["so2", "0 to 2000 µg/m³"],
            ["nh3", "0 to 2000 µg/m³"],
            ["temperature", "−30 to 60 °C"],
            ["humidity", "0 to 100 %"],
            ["wind_speed", "0 to 150 km/h"],
            ["wind_deg", "0 to 360 degrees"],
            ["pressure", "850 to 1100 hPa"],
        ],
    )
    caption(doc, "Table 10: Validation ranges applied before insert.")
    body(
        doc,
        "The bounds are wide on purpose so genuine smog extremes are not treated as errors. "
        "One bad hour drops that row; it does not abort the whole batch unless I set that flag."
    )
    heading(doc, "9.2 Hourly grid and short-gap interpolation", 2)
    body(
        doc,
        "I reindex onto a strict hourly calendar before any lag or target shift. The raw history "
        "has about **3,700 missing hours** in 415 outages. On the gappy frame a 24-hour shift "
        "spanned 24 hours for only 87% of rows and up to 264 hours for the rest. Outages of at "
        "most six hours, end to end, I interpolate linearly. Longer gaps stay missing and drop "
        "later. Integer-typed columns are rounded immediately after the grid step."
    )
    heading(doc, "9.3 Feature transforms", 2)
    numbered(
        doc,
        [
            "Time features: hour, day_of_week, month, is_weekend, hour_sin, hour_cos, month_sin, month_cos.",
            "Smog flag: is_smog_season for October–January.",
            "log1p of clipped-at-zero co, no, no2, o3, so2, pm2_5, pm10, nh3 (raw columns replaced in place).",
            "AQI lags at 1, 3, 6, 24 and 168 hours.",
            "Rolling AQI mean, std, min, max over 3, 6 and 24 hours.",
            "Change rates: aqi_change_rate_1h and aqi_change_rate_24h.",
            "Weather interactions: wind_speed × pm2_5 and humidity × pm2_5.",
            "Targets: aqi_target_{24,48,72}h and aqi_delta_{24,48,72}h.",
            "Training mode: drop any row with a missing value. Inference mode: drop missing values only on feature columns so the latest hours (unknown future targets) remain.",
        ],
    )
    heading(doc, "9.4 Scaling", 2)
    body(
        doc,
        "Ridge is a pipeline of StandardScaler plus Ridge. LSTM and GRU fit StandardScaler on "
        "training features and then transform train, validation and test. RandomForest, XGBoost "
        "and LightGBM are not scaled. Prophet uses timestamp and AQI only."
    )
    heading(doc, "9.5 Train / validation / test split", 2)
    body(
        doc,
        "I used a chronological split that snaps validation and test starts to 1 June. The run "
        "printed: train 2020-12-05→2024-06-01 | val 2024-06-01→2025-06-01 | "
        "test 2025-06-01→2026-08-08, with **27,947 / 8,075 / 10,078** rows. Optuna used "
        "TimeSeriesSplit with 5 folds on the training fold. Random state 42. 15 trials."
    )
    heading(doc, "9.6 What I did not do in the first training run", 2)
    body(
        doc,
        "That notebook did not start training on 4 April 2025. It also did not search delta "
        "shrinkage (that search exists in the production training loop). The raw snapshot had "
        "no duplicate timestamps, so duplicate removal was not needed there."
    )

    heading(doc, "10. Exploratory Data Analysis", 1)
    body(
        doc,
        "I plotted the raw snapshot using the same recipes as my EDA notebook (univariate AQI, "
        "pollutant histograms, AQI over time, AQI versus weather, correlation heatmap, ACF/PACF, "
        "seasonal decomposition, ADF, and smog versus normal). The notebook in git has no saved "
        "plot outputs, so I regenerated the figures below from the parquet."
    )

    heading(doc, "10.1 AQI distribution", 2)
    add_picture(doc, "eda_aqi_dist.png", 6.3)
    caption(doc, "Figure 1: Distribution of aqi (histogram with KDE) and aqi boxplot on the raw snapshot.")
    body(
        doc,
        "AQI is right-skewed (**skewness 1.91**). The mean (247.8) sits well above the median "
        "(177). The boxplot has a long upper tail: **5,837** hours sit above the IQR fence of 479. "
        "The maximum in both smog and normal groups is 1000."
    )

    heading(doc, "10.2 Pollutant distributions", 2)
    add_picture(doc, "eda_pollutant_hist.png", 6.3)
    caption(doc, "Figure 2: Histograms of pm2_5, pm10, co, no, no2, o3, so2 and nh3 on the raw snapshot.")
    body(
        doc,
        "All eight pollutant series are right-skewed. Ranked skewness is in Section 8. That is "
        "why I log-transform them before modelling."
    )

    heading(doc, "10.3 AQI over time", 2)
    add_picture(doc, "eda_aqi_time.png", 6.3)
    caption(doc, "Figure 3: aqi over time for Lahore on the raw snapshot (2020-11-27 to 2026-08-14).")
    body(
        doc,
        "Winter peaks repeat every year, and the typical level drifts down later — yearly means "
        "in Table 7 fall from 438.8 in incomplete 2020 to 166.5 in incomplete 2026. That multi-year "
        "downward trend is one reason I train on **deltas** rather than raw level."
    )

    heading(doc, "10.4 AQI versus weather", 2)
    add_picture(doc, "eda_aqi_weather.png", 6.3)
    caption(doc, "Figure 4: Scatter of aqi versus temperature, humidity, wind_speed and pressure.")
    body(
        doc,
        "Temperature correlates at **−0.55**, pressure at **+0.51**, humidity at +0.28 and wind "
        "speed at −0.24 (Table 9). Cooler, higher-pressure hours tend to sit with worse AQI in "
        "this city. I treat those as associations, not a causal story, though wind dispersing "
        "pollution and humidity affecting secondary particles are the usual physical reading."
    )

    heading(doc, "10.5 Correlation heatmap", 2)
    add_picture(doc, "eda_corr.png", 5.6)
    caption(doc, "Figure 5: Correlation heatmap of numeric raw columns (month excluded).")
    body(
        doc,
        "AQI is almost collinear with **pm2_5 (0.986)** and very close to pm10 (0.958), which is "
        "expected because I compute AQI from PM2.5. OpenWeather’s 1–5 category only correlates "
        "at 0.48, which is why I did not use it as the target."
    )

    heading(doc, "10.6 Autocorrelation", 2)
    add_picture(doc, "eda_acf_pacf.png", 6.3)
    caption(doc, "Figure 6: ACF of aqi up to 168 hourly lags and PACF of aqi up to 72 hourly lags.")
    body(
        doc,
        "Short-lag autocorrelation is strong, and the ACF still has structure at daily (24h) and "
        "weekly (168h) ranges. That is why I kept lags {1, 3, 6, 24, 168} and rolling windows "
        "{3, 6, 24}."
    )

    heading(doc, "10.7 Seasonal decomposition", 2)
    add_picture(doc, "eda_decompose.png", 6.0)
    caption(doc, "Figure 7: Additive seasonal_decompose of hourly aqi with period=24.")
    body(
        doc,
        "An additive 24-hour decomposition shows a slow trend, a repeating daily seasonal piece, "
        "and a residual. I interpolated to hourly frequency before decompose so the period is a "
        "true 24 hours."
    )

    heading(doc, "10.8 Smog season versus normal season", 2)
    add_picture(doc, "eda_smog.png", 5.2)
    caption(doc, "Figure 8: Boxplot of aqi for Normal versus Smog (Oct–Jan) on the raw snapshot.")
    body(
        doc,
        "Smog-season hours are much worse: mean **399** (n = 14,013) versus **182** (n = 32,406). "
        "I keep `is_smog_season` both as a model feature and as a way to slice the data."
    )

    heading(doc, "11. Feature and Variable Analysis", 1)
    body(
        doc,
        "The first notebook used **42 input columns**, after dropping timestamp, hour, month, "
        "the OpenWeather 1–5 field, and all six target columns from the 52-column feature table. "
        "The 42 names are:"
    )
    bullets(
        doc,
        [
            "Pollutants after log1p (replaced in place): co, no, no2, o3, so2, pm2_5, pm10, nh3.",
            "Current aqi (untransformed).",
            "Weather: temperature, humidity, wind_speed, wind_deg, pressure.",
            "Time encodings kept as features: day_of_week, is_weekend, hour_sin, hour_cos, month_sin, month_cos, is_smog_season.",
            "AQI lags: aqi_lag_1h, aqi_lag_3h, aqi_lag_6h, aqi_lag_24h, aqi_lag_168h.",
            "Rolling aqi: mean/std/min/max for windows 3h, 6h and 24h (12 columns).",
            "Change rates: aqi_change_rate_1h, aqi_change_rate_24h.",
            "Interactions: wind_pollutant_interaction, humidity_pollutant_interaction.",
        ],
    )
    body(
        doc,
        "That list has 42 names. The notebook printed `n_features: 42` but did not dump the "
        "names as a separate list; I reconstructed them from the feature builder plus the drop list."
    )
    body(
        doc,
        "Tabular models, LSTM and GRU train on `aqi_delta_{h}h`. Absolute metrics use "
        "`aqi_target_{h}h`. Persistence predicts the current `aqi` column. Prophet is fit on "
        "timestamp and AQI and predicts at test time plus h hours."
    )

    heading(doc, "12. Methodology", 1)
    body(
        doc,
        "All families except Prophet predict the AQI change over the horizon. I then add that "
        "delta to current AQI. Metrics are RMSE, MAE and R² on the absolute future AQI. "
        "Persistence is scored the same way against current AQI."
    )
    body(
        doc,
        "Registration rule: RMSE must be **lower** than persistence, MAE **lower**, and R² "
        "**higher**. Among eligible models I sort by RMSE, then MAE, then −R². If none are "
        "eligible, I register nothing."
    )
    body(
        doc,
        "That first run logged into Hopsworks project https://eu-west.cloud.hopsworks.ai:443/p/41103 "
        "and read `aqi_features` v3. TensorFlow saw a GPU; XGBoost used CUDA."
    )

    heading(doc, "13. Approaches Considered and Used", 1)
    heading(doc, "13.1 Persistence baseline", 2)
    body(
        doc,
        "No-change forecast: future AQI equals current AQI. No extra features and no training. "
        "I score it with RMSE, MAE and R² on the held-out test fold (Tables 11–13 in the "
        "per-model section, and Tables 21–23 in the full comparison). In the first notebook this "
        "baseline **won every horizon**, so nothing else was registered."
    )
    heading(doc, "13.2 Optuna-tuned tabular regression on deltas", 2)
    body(
        doc,
        "Ridge, RandomForest, XGBoost and LightGBM predict the horizon delta from the 42 features. "
        "I tuned each with Optuna (15 trials, time-series CV RMSE) and refit on all training rows. "
        "None beat persistence on all three metrics in that run."
    )
    heading(doc, "13.3 Mean ensemble of top-3 tabular models", 2)
    body(
        doc,
        "I averaged the reconstructed absolute predictions of the three tabular models with "
        "lowest test RMSE. That ensemble was not eligible under the persistence gate either."
    )
    heading(doc, "13.4 Sequence models (LSTM and GRU)", 2)
    body(
        doc,
        "LSTM and GRU see 24-hour windows of the 42 scaled features and predict the horizon delta. "
        "Architecture: Input (24, n_feat), LSTM/GRU 64 with return_sequences, Dropout 0.2, "
        "LSTM/GRU 32, Dropout 0.2, Dense 16 ReLU, Dense 1. Adam at 1e-3, MSE loss, 100 epochs, "
        "batch 128, early stopping on val_loss with patience 8. Both had RMSE above 219 and "
        "**negative R²** on every horizon. I did not register them."
    )
    heading(doc, "13.5 Prophet on absolute AQI", 2)
    body(
        doc,
        "Prophet models the AQI level directly. I fit it on train timestamp and AQI with daily "
        "and weekly seasonality on, yearly seasonality off. Future dates are test timestamp plus "
        "the horizon. RMSE was above **2150** on every horizon. Not registered."
    )
    heading(doc, "13.6 Production training window", 2)
    body(
        doc,
        "After the OpenWeather break on 4 April 2025 I started training from that date only. "
        "That is a different experiment from the first notebook. Results for that window are "
        "in Section 20, not mixed into Section 17."
    )
    heading(doc, "13.7 Classification of OpenWeather 1–5", 2)
    body(
        doc,
        "OpenWeather also publishes a 1–5 air-quality class. I ran that as a separate experiment: "
        "predict tomorrow’s class from the same features, with persistence meaning “today’s class.” "
        "Models were ordinal logistic regression, Random Forest, XGBoost, LightGBM, a majority-vote "
        "ensemble, LSTM and GRU (5-way softmax). A model had to beat persistence on accuracy, "
        "macro F1, RMSE, MAE and R² together (the last three on the 1–5 labels). Results are in "
        "Section 21. I did not adopt this path."
    )

    heading(doc, "14. Models Used", 1)
    body(
        doc,
        "Numbers below are the full-precision test scores from the first training notebook. "
        "One-line prints in that notebook used fewer decimals; the tables keep the full comparison."
    )

    heading(doc, "14.1 Persistence Baseline", 2)
    body(doc, "No-change forecast, used as the registration gate. Input is current AQI only. No training.")
    add_table(
        doc,
        ["Horizon", "RMSE", "MAE", "R2"],
        [
            ["24", "66.730405", "36.111431", "0.745204"],
            ["48", "84.796609", "45.775352", "0.588437"],
            ["72", "92.456124", "51.518059", "0.510480"],
        ],
    )
    caption(doc, "Table 11: Persistence Baseline test metrics.")
    body(
        doc,
        "On that test window, predicting no change already explains a large share of variance "
        "at 24h (**R² 0.745**) and still more than half at 72h (R² 0.510). RMSE rises with "
        "horizon. Under the gate I wrote, this is the selected forecast from that notebook."
    )

    heading(doc, "14.2 Ridge", 2)
    body(
        doc,
        "Linear model of the AQI delta with L2 penalty, after StandardScaler. I searched "
        "alpha in [0.01, 100] (log-uniform, 15 trials) and then fit StandardScaler + Ridge "
        "on the training fold."
    )
    add_table(
        doc,
        ["Horizon", "RMSE", "MAE", "R2"],
        [
            ["24", "87.584091", "63.751444", "0.561069"],
            ["48", "101.382767", "73.917026", "0.411688"],
            ["72", "104.353359", "77.654450", "0.376391"],
        ],
    )
    caption(doc, "Table 12: Ridge test metrics.")
    body(
        doc,
        "Ridge RMSE and MAE are higher than persistence at every horizon, and R² is lower. "
        "I did not select it."
    )

    heading(doc, "14.3 RandomForest", 2)
    body(
        doc,
        "Bagged trees on the AQI delta. Optuna over n_estimators 100–250 (step 50), max_depth "
        "in {8, 12, 16}, min_samples_leaf 1–8, with max_samples=0.8, 15 trials."
    )
    add_table(
        doc,
        ["Horizon", "RMSE", "MAE", "R2"],
        [
            ["24", "84.914023", "59.164036", "0.587424"],
            ["48", "118.050540", "86.744712", "0.202344"],
            ["72", "116.121902", "85.821522", "0.227804"],
        ],
    )
    caption(doc, "Table 13: RandomForest test metrics.")
    body(
        doc,
        "Worse than persistence on RMSE, MAE and R² at every horizon. Not selected."
    )

    heading(doc, "14.4 XGBoost", 2)
    body(
        doc,
        "Gradient-boosted trees on the AQI delta. Optuna over learning_rate 0.01–0.3 (log-uniform), "
        "max_depth 3–9, n_estimators 100–400 (step 50), subsample and colsample_bytree 0.6–1.0. "
        "This run used hist + CUDA."
    )
    add_table(
        doc,
        ["Horizon", "RMSE", "MAE", "R2"],
        [
            ["24", "74.364967", "50.019281", "0.683567"],
            ["48", "97.213262", "69.199877", "0.459083"],
            ["72", "111.990073", "82.949948", "0.281779"],
        ],
    )
    caption(doc, "Table 14: XGBoost test metrics.")
    body(
        doc,
        "Among learned models, XGBoost had the lowest 24h and 48h RMSE, but both RMSE and MAE "
        "remain above persistence and R² remains below it. Closest at 24h (**74.36 vs 66.73**), "
        "still not eligible. Not selected."
    )

    heading(doc, "14.5 LightGBM", 2)
    body(
        doc,
        "Gradient-boosted trees on the AQI delta. Optuna over learning_rate 0.01–0.3, "
        "num_leaves 16–96, n_estimators 100–400 (step 50), subsample and colsample_bytree 0.6–1.0."
    )
    add_table(
        doc,
        ["Horizon", "RMSE", "MAE", "R2"],
        [
            ["24", "76.456676", "52.289149", "0.665515"],
            ["48", "101.157242", "72.808352", "0.414302"],
            ["72", "112.211034", "81.688058", "0.278942"],
        ],
    )
    caption(doc, "Table 15: LightGBM test metrics.")
    body(
        doc,
        "Close to XGBoost at 24h, still behind persistence on all three metrics. Not selected."
    )

    heading(doc, "14.6 Ensemble_top3_tabular", 2)
    body(
        doc,
        "Mean of the three best tabular absolute reconstructions by test RMSE. No extra fit; "
        "just an average of models already trained."
    )
    add_table(
        doc,
        ["Horizon", "RMSE", "MAE", "R2"],
        [
            ["24", "77.085201", "53.331851", "0.659993"],
            ["48", "98.112376", "70.283343", "0.449031"],
            ["72", "107.619632", "79.190920", "0.336742"],
        ],
    )
    caption(doc, "Table 16: Ensemble_top3_tabular test metrics.")
    body(
        doc,
        "The ensemble did not beat the best single tabular model at 24h or 48h and did not beat "
        "persistence. Not selected."
    )

    heading(doc, "14.7 LSTM", 2)
    body(
        doc,
        "Recurrent net on 24-hour windows of scaled features, predicting the delta. Training "
        "setup is in Section 13.4."
    )
    add_table(
        doc,
        ["Horizon", "RMSE", "MAE", "R2"],
        [
            ["24", "232.032646", "156.863391", "-2.075345"],
            ["48", "234.142256", "152.548810", "-2.132563"],
            ["72", "235.968967", "160.365461", "-2.182948"],
        ],
    )
    caption(doc, "Table 17: LSTM test metrics.")
    body(
        doc,
        "RMSE more than three times persistence at 24h; R² is negative, so the forecast is worse "
        "than predicting the test-set mean. Not selected."
    )

    heading(doc, "14.8 GRU", 2)
    body(doc, "Same sequence setup as LSTM, with GRU layers.")
    add_table(
        doc,
        ["Horizon", "RMSE", "MAE", "R2"],
        [
            ["24", "219.023846", "146.531186", "-1.740176"],
            ["48", "246.401500", "157.766790", "-2.469181"],
            ["72", "239.031295", "159.196560", "-2.266098"],
        ],
    )
    caption(doc, "Table 18: GRU test metrics.")
    body(
        doc,
        "Slightly better than LSTM at 24h, still far worse than persistence, with negative R². "
        "Not selected."
    )

    heading(doc, "14.9 Prophet", 2)
    body(doc, "Additive time-series model of absolute AQI. Inputs are train timestamp and train AQI only.")
    add_table(
        doc,
        ["Horizon", "RMSE", "MAE", "R2"],
        [
            ["24", "2167.324472", "1736.633538", "-267.777676"],
            ["48", "2159.443832", "1727.409530", "-265.909192"],
            ["72", "2150.783604", "1715.485984", "-263.906508"],
        ],
    )
    caption(doc, "Table 19: Prophet test metrics.")
    body(
        doc,
        "RMSE above 2150 AQI points on every horizon. Yearly seasonality was off. I did not "
        "dig further into fitted parameters after that. Not selected."
    )

    heading(doc, "15. Model Training", 1)
    body(
        doc,
        "Per horizon I computed persistence first, then trained Prophet, then tuned Ridge, "
        "RandomForest, XGBoost and LightGBM on training deltas, then trained LSTM and GRU, "
        "then built the top-3 tabular ensemble, then printed a table sorted by RMSE and applied "
        "the registration gate. If none passed, I registered nothing."
    )
    body(
        doc,
        "The notebook can write per-horizon comparison CSVs. The numbers in this report come "
        "from the saved cell output in that notebook, not from those CSVs."
    )

    heading(doc, "16. Model Evaluation", 1)
    body(
        doc,
        "On the regression path I used RMSE, MAE and R² on absolute future AQI (0–500). "
        "The classification path in Section 21 also used accuracy and macro F1 on the 1–5 labels."
    )
    body(
        doc,
        "The gate is: RMSE smaller than persistence, MAE smaller, R² larger. The printed outcome "
        "for each horizon was that **no model beat persistence on all three metrics**, so nothing "
        "was registered:"
    )
    add_table(
        doc,
        ["horizon", "registered", "winner"],
        [
            ["24", "False", "None"],
            ["48", "False", "None"],
            ["72", "False", "None"],
        ],
    )
    caption(doc, "Table 20: Registry summary from the first training notebook.")

    heading(doc, "17. Results", 1)
    heading(doc, "17.1 Complete comparison (24 hours)", 2)
    add_table(doc, ["Model", "RMSE", "MAE", "R2"], metric_rows(AT8_24))
    caption(doc, "Table 21: Test comparison for horizon 24h.")
    add_picture(doc, "at8_rmse_24h.png", 6.3)
    caption(
        doc,
        "Figure 9: Test RMSE by model at 24h using the values in Table 21. Persistence is the top bar after sorting by RMSE. Prophet is omitted from this bar chart because RMSE 2167.324472 would flatten the other bars; Prophet remains in Table 21.",
    )
    body(
        doc,
        "Every learned model has higher RMSE and higher MAE than persistence, and lower R². "
        "XGBoost is the closest learned model (**RMSE 74.36 versus 66.73**)."
    )

    heading(doc, "17.2 Complete comparison (48 hours)", 2)
    add_table(doc, ["Model", "RMSE", "MAE", "R2"], metric_rows(AT8_48))
    caption(doc, "Table 22: Test comparison for horizon 48h.")
    add_picture(doc, "at8_rmse_48h.png", 6.3)
    caption(
        doc,
        "Figure 10: Test RMSE by model at 48h using the values in Table 22. Prophet omitted from bars for the same scale reason as Figure 9.",
    )
    body(
        doc,
        "Persistence remains first (**RMSE 84.80**). XGBoost is again the closest learned model "
        "(97.21). LSTM, GRU and Prophet remain far worse."
    )

    heading(doc, "17.3 Complete comparison (72 hours)", 2)
    add_table(doc, ["Model", "RMSE", "MAE", "R2"], metric_rows(AT8_72))
    caption(doc, "Table 23: Test comparison for horizon 72h.")
    add_picture(doc, "at8_rmse_72h.png", 6.3)
    caption(
        doc,
        "Figure 11: Test RMSE by model at 72h using the values in Table 23. Prophet omitted from bars for the same scale reason as Figure 9.",
    )
    body(
        doc,
        "Persistence RMSE **92.46** is still lowest. Among learned models, Ridge has the lowest "
        "72h RMSE (104.35), still above persistence."
    )

    heading(doc, "18. Model Comparison", 1)
    body(
        doc,
        "Table 24 puts persistence beside the best learned model by RMSE at each horizon. "
        "“Best learned” means lowest RMSE among non-persistence rows. It is not a registration winner."
    )
    add_table(
        doc,
        ["Horizon", "Persistence RMSE", "Best learned model by RMSE", "That model RMSE", "Selected"],
        [
            ["24", "66.730405", "XGBoost", "74.364967", "No — persistence wins"],
            ["48", "84.796609", "XGBoost", "97.213262", "No — persistence wins"],
            ["72", "92.456124", "Ridge", "104.353359", "No — persistence wins"],
        ],
    )
    caption(doc, "Table 24: Persistence versus lowest-RMSE learned model.")

    heading(doc, "19. Best Model on the First Training Notebook", 1)
    body(
        doc,
        "On the first notebook (full history, feature group v3) I selected **no learned model**. "
        "Persistence is the only forecast that satisfies the three-metric rule on that split. "
        "Exact persistence numbers are in Table 11. Registered was False and winner was None "
        "for 24, 48 and 72 hours."
    )
    body(
        doc,
        "That is not my production choice. It is the result of that one experiment. The later "
        "regression window is in Section 20. Classification is in Section 21. I state the "
        "final decision in Section 22."
    )

    heading(doc, "20. Production Training Window", 1)
    body(
        doc,
        "This section is a later experiment. Training across the 4 April 2025 break is what made "
        "every model lose to persistence on the 2025–26 test period, so I now start training on "
        "that date. Selected models for that window:"
    )
    add_table(
        doc,
        ["Horizon", "Selected model", "RMSE", "MAE", "R²"],
        [
            ["24 h (day 1)", "LightGBM (L1)", "28.88", "22.63", "0.064"],
            ["48 h (day 2)", "Random Forest", "30.54", "24.89", "−0.055"],
            ["72 h (day 3)", "Ensemble (top-3 tabular)", "31.65", "26.07", "−0.098"],
        ],
    )
    caption(
        doc,
        "Table 25: Selected models on the post-4 April 2025 training window. Not from the first notebook.",
    )
    add_table(
        doc,
        ["Horizon", "Persistence RMSE", "Selected-model RMSE"],
        [
            ["24 h", "34.84", "28.88"],
            ["48 h", "40.94", "30.54"],
            ["72 h", "40.30", "31.65"],
        ],
    )
    caption(
        doc,
        "Table 26: Persistence versus selected RMSE on the post-4 April 2025 window.",
    )
    body(
        doc,
        "After mentor review I treated this window as the production choice, with registry names "
        "`aqi_forecaster_24h`, `aqi_forecaster_48h` and `aqi_forecaster_72h`. I do not replace "
        "Tables 21–23 with Tables 25–26. The two experiments differ in training start date and "
        "in feature-group version (v3 then, v4 now)."
    )

    heading(doc, "20.1 SHAP analysis of registered models", 2)
    body(
        doc,
        "The first notebook has a SHAP block, but it only runs when a model is registered. "
        "That run registered nothing, so it produced no SHAP plots."
    )
    body(
        doc,
        "I therefore computed SHAP myself on the models currently cached from the registry, "
        "using TreeExplainer — the same method the training pipeline uses. Features come from "
        "the local raw parquet restricted to **4 April 2025** onwards. The chronological split "
        "fell back to 70/15/15 on that window: **6,956 / 1,491 / 1,491** rows (test from "
        "29 May 2026 11:00 to 11 August 2026 13:00). I sampled 400 test rows for gradient-boosted "
        "trees and **150** for RandomForest (a 2,000-row forest run timed out)."
    )
    body(
        doc,
        "The cached payloads are named **XGBoost** (24h), **RandomForest** (48h) and "
        "**Ensemble_top3_tabular** (72h members XGBoost_L1, RandomForest, LightGBM). Those names "
        "come from the saved models, not from Table 25. SHAP explains the predicted AQI delta "
        "(target_type delta, shrinkage 1.0)."
    )

    heading(doc, "20.1.1 24-hour registered model (XGBoost)", 3)
    add_table(
        doc,
        ["Rank", "Feature", "Mean |SHAP|"],
        [
            ["1", "pressure", "5.7132248878479"],
            ["2", "aqi_change_rate_24h", "3.0580039024353027"],
            ["3", "month_cos", "2.942716598510742"],
            ["4", "pm10", "2.813547134399414"],
            ["5", "temperature", "2.3712317943573"],
            ["6", "humidity_pollutant_interaction", "2.0532147884368896"],
            ["7", "wind_deg", "1.9747484922409058"],
            ["8", "day_of_week", "1.7397871017456055"],
            ["9", "aqi_roll_min_24h", "1.629075527191162"],
            ["10", "aqi_roll_min_6h", "1.5423707962036133"],
        ],
    )
    caption(doc, "Table 27: Top 10 mean |SHAP| features for the 24h XGBoost model (400 test rows).")
    add_picture(doc, "shap_24h_beeswarm.png", 6.2)
    caption(doc, "Figure 12: SHAP summary (beeswarm) for the 24h XGBoost model.")
    body(
        doc,
        "Each point is one sampled test hour. Colour is the feature value (high versus low). "
        "Horizontal position is that feature’s contribution to the predicted delta. "
        "**pressure** has the largest mean |SHAP| (5.71), then aqi_change_rate_24h (3.06) and "
        "month_cos (2.94)."
    )
    add_picture(doc, "shap_24h_bar.png", 5.8)
    caption(doc, "Figure 13: Mean |SHAP| bar chart for the 24h XGBoost model (top 15 features).")

    heading(doc, "20.1.2 48-hour registered model (RandomForest)", 3)
    add_table(
        doc,
        ["Rank", "Feature", "Mean |SHAP|"],
        [
            ["1", "pm2_5", "8.678609856675047"],
            ["2", "pressure", "5.349003206462411"],
            ["3", "month_cos", "4.57826334528948"],
            ["4", "aqi_change_rate_24h", "3.6110002145874653"],
            ["5", "aqi_lag_168h", "3.0203817288733674"],
            ["6", "day_of_week", "2.8400790645516274"],
            ["7", "month_sin", "2.4771160689993983"],
            ["8", "temperature", "2.2241804783036776"],
            ["9", "aqi_roll_max_24h", "1.9557856727520853"],
            ["10", "nh3", "1.48097618134355"],
        ],
    )
    caption(doc, "Table 28: Top 10 mean |SHAP| features for the 48h RandomForest model (150 test rows).")
    add_picture(doc, "shap_48h_beeswarm.png", 6.2)
    caption(doc, "Figure 14: SHAP summary (beeswarm) for the 48h RandomForest model.")
    body(
        doc,
        "**pm2_5** has the largest mean |SHAP| (8.68), then pressure (5.35) and month_cos (4.58). "
        "This ranking is for 150 sampled test rows."
    )
    add_picture(doc, "shap_48h_bar.png", 5.8)
    caption(doc, "Figure 15: Mean |SHAP| bar chart for the 48h RandomForest model (top 15 features).")

    heading(doc, "20.1.3 72-hour registered ensemble members", 3)
    body(
        doc,
        "The 72h model is a mean ensemble. I ran TreeExplainer on each member because there is "
        "no single SHAP plot for the average."
    )
    add_table(
        doc,
        ["Member", "Top feature", "Mean |SHAP|", "SHAP rows"],
        [
            ["XGBoost_L1", "month_cos", "6.110957622528076", "400"],
            ["RandomForest", "pm2_5", "14.038097403347498", "150"],
            ["LightGBM", "month_cos", "5.241344214210389", "400"],
        ],
    )
    caption(doc, "Table 29: Highest mean |SHAP| feature per 72h ensemble member.")
    add_picture(doc, "shap_72h_XGBoost_L1_beeswarm.png", 6.2)
    caption(doc, "Figure 16: SHAP summary for 72h ensemble member XGBoost_L1.")
    add_picture(doc, "shap_72h_XGBoost_L1_bar.png", 5.8)
    caption(doc, "Figure 17: Mean |SHAP| bars for 72h XGBoost_L1.")
    add_picture(doc, "shap_72h_RandomForest_beeswarm.png", 6.2)
    caption(doc, "Figure 18: SHAP summary for 72h ensemble member RandomForest.")
    add_picture(doc, "shap_72h_RandomForest_bar.png", 5.8)
    caption(doc, "Figure 19: Mean |SHAP| bars for 72h RandomForest.")
    add_picture(doc, "shap_72h_LightGBM_beeswarm.png", 6.2)
    caption(doc, "Figure 20: SHAP summary for 72h ensemble member LightGBM.")
    add_picture(doc, "shap_72h_LightGBM_bar.png", 5.8)
    caption(doc, "Figure 21: Mean |SHAP| bars for 72h LightGBM.")
    body(
        doc,
        "On these samples, **month_cos** is the top mean |SHAP| feature for XGBoost_L1 and LightGBM. "
        "**pm2_5** is the top feature for the RandomForest member (14.04). `is_smog_season` "
        "appears in the XGBoost_L1 and RandomForest top-10 lists. I am not reading these rankings "
        "as a causal story."
    )

    heading(doc, "21. Classification Experiment", 1)
    body(
        doc,
        "I also tried classification. The target was future OpenWeather class "
        "(openweather_aqi_category, values 1–5) at 24, 48 and 72 hours, using the same feature "
        "matrix as production. Persistence here means “tomorrow’s class equals today’s class.” "
        "I trained ordinal logistic regression, Random Forest, XGBoost, LightGBM, a majority-vote "
        "ensemble, LSTM and GRU with a 5-way softmax. A model counted as a winner only if it beat "
        "persistence on **accuracy, macro F1, RMSE, MAE and R²** together. RMSE, MAE and R² in "
        "this section are on the 1–5 labels, not on 0–500 AQI. Nothing from this experiment "
        "went into the registry."
    )

    heading(doc, "21.1 Full-history classification", 2)
    body(
        doc,
        "On the full archive (December 2020–August 2026, about 39,060 hours) persistence accuracy "
        "was **0.633 / 0.578 / 0.548** at 24 / 48 / 72 hours. Learned models sat around **0.31–0.38**. "
        "No model passed the five-metric gate. Persistence also has the lowest RMSE at every horizon "
        "(Table 30)."
    )
    add_table(
        doc,
        ["H", "Model", "RMSE", "MAE", "R²"],
        [
            ["24", "Persistence", "0.729", "0.420", "0.412"],
            ["24", "Ordinal logistic", "1.461", "1.131", "−1.358"],
            ["24", "Random Forest", "1.307", "0.964", "−0.887"],
            ["24", "XGBoost", "1.243", "0.905", "−0.707"],
            ["24", "LightGBM", "1.322", "0.972", "−0.930"],
            ["24", "Ensemble", "1.286", "0.943", "−0.827"],
            ["24", "LSTM", "1.494", "1.155", "−1.462"],
            ["24", "GRU", "1.472", "1.127", "−1.390"],
            ["48", "Persistence", "0.816", "0.500", "0.264"],
            ["48", "Ordinal logistic", "1.512", "1.178", "−1.525"],
            ["48", "Random Forest", "1.493", "1.150", "−1.462"],
            ["48", "XGBoost", "1.353", "1.001", "−1.022"],
            ["48", "LightGBM", "1.334", "0.977", "−0.965"],
            ["48", "Ensemble", "1.392", "1.041", "−1.141"],
            ["48", "LSTM", "1.512", "1.178", "−1.524"],
            ["48", "GRU", "1.503", "1.161", "−1.493"],
            ["72", "Persistence", "0.863", "0.546", "0.178"],
            ["72", "Ordinal logistic", "1.512", "1.176", "−1.525"],
            ["72", "Random Forest", "1.510", "1.173", "−1.519"],
            ["72", "XGBoost", "1.402", "1.049", "−1.171"],
            ["72", "LightGBM", "1.397", "1.050", "−1.157"],
            ["72", "Ensemble", "1.440", "1.092", "−1.291"],
            ["72", "LSTM", "1.521", "1.186", "−1.552"],
            ["72", "GRU", "1.513", "1.173", "−1.525"],
        ],
    )
    caption(doc, "Table 30: Classification on full history. Scores are on OpenWeather classes 1–5. Persistence wins every horizon.")

    heading(doc, "21.2 Post-break classification", 2)
    body(
        doc,
        "On the April 2025–August 2026 window (about 9,866 hours) the race was closer. Only "
        "**72-hour Random Forest** passed all five metrics: accuracy **0.568** versus persistence "
        "**0.489**, RMSE **0.746** versus **0.895**. Forty-eight-hour GRU looked strong on accuracy "
        "(0.568 versus 0.486) but lost on macro F1 (0.292 versus 0.298). A walk-forward check "
        "(four folds, Random Forest versus persistence) passed the five-metric test in **1 of 12** "
        "fold×horizon slots, so the 72-hour win is not stable."
    )
    add_table(
        doc,
        ["H", "Model", "RMSE", "MAE", "R²"],
        [
            ["24", "Persistence", "0.826", "0.520", "−0.558"],
            ["24", "Ordinal logistic", "0.753", "0.472", "−0.297"],
            ["24", "Random Forest", "0.780", "0.497", "−0.391"],
            ["24", "XGBoost", "0.821", "0.539", "−0.541"],
            ["24", "LightGBM", "0.817", "0.543", "−0.527"],
            ["24", "Ensemble", "0.813", "0.532", "−0.510"],
            ["24", "LSTM", "0.778", "0.495", "−0.379"],
            ["24", "GRU", "0.800", "0.526", "−0.455"],
            ["48", "Persistence", "0.883", "0.602", "−0.789"],
            ["48", "Ordinal logistic", "0.796", "0.524", "−0.456"],
            ["48", "Random Forest", "0.812", "0.534", "−0.513"],
            ["48", "XGBoost", "0.802", "0.540", "−0.476"],
            ["48", "LightGBM", "0.829", "0.560", "−0.578"],
            ["48", "Ensemble", "0.807", "0.539", "−0.496"],
            ["48", "LSTM", "0.797", "0.525", "−0.488"],
            ["48", "GRU", "0.758", "0.480", "−0.346"],
            ["72", "Persistence", "0.895", "0.607", "−0.888"],
            ["72", "Ordinal logistic", "0.776", "0.502", "−0.420"],
            ["72", "Random Forest *", "0.746", "0.474", "−0.311"],
            ["72", "XGBoost", "0.845", "0.570", "−0.683"],
            ["72", "LightGBM", "0.868", "0.594", "−0.777"],
            ["72", "Ensemble", "0.822", "0.548", "−0.592"],
            ["72", "LSTM", "0.799", "0.530", "−0.495"],
            ["72", "GRU", "0.767", "0.489", "−0.380"],
        ],
    )
    caption(
        doc,
        "Table 31: Classification on the post-4 April 2025 window. *Only 72h Random Forest beat persistence on all five metrics (accuracy and macro F1 included).",
    )
    add_table(
        doc,
        ["Setting", "Horizon", "Persistence accuracy", "Best learned (note)"],
        [
            ["Full history", "24 h", "0.633", "Learned models ~0.31–0.38; none passed the gate"],
            ["Full history", "48 h", "0.578", "Learned models ~0.31–0.38; none passed the gate"],
            ["Full history", "72 h", "0.548", "Learned models ~0.31–0.38; none passed the gate"],
            ["Post-break", "24 h", "—", "No five-metric winner"],
            ["Post-break", "48 h", "0.486", "GRU accuracy 0.568, but macro F1 0.292 vs 0.298"],
            ["Post-break", "72 h", "0.489", "Random Forest accuracy 0.568 (passed all five; not stable in walk-forward)"],
        ],
    )
    caption(doc, "Table 32: Classification accuracy summary. Macro F1 is stated where it decided the gate.")
    body(
        doc,
        "Classifying 1–5 is a different product from forecasting 0–500 AQI. It does not beat "
        "“today’s class” in a reliable way. I did not replace the continuous regression path."
    )

    heading(doc, "22. Final Approach Selection", 1)
    body(
        doc,
        "I compared regression and classification, each on the full archive and on the post-4 April "
        "2025 window. Table 33 records what each experiment measured, which model won, and whether "
        "I kept it. The decision at the end of that comparison is the production system."
    )
    add_table(
        doc,
        ["Experiment", "What I measured", "What won", "Kept?"],
        [
            ["Regression, full history (first notebook)", "RMSE, MAE, R² on 0–500 AQI", "Persistence at 24 / 48 / 72 h (RMSE 66.73 / 84.80 / 92.46). Closest learned 24h model: XGBoost 74.36.", "No"],
            ["Regression, from 4 April 2025", "RMSE, MAE, R² on 0–500 AQI", "LightGBM 24h (RMSE 28.88), Random Forest 48h (30.54), top-3 ensemble 72h (31.65). All three beat persistence.", "Yes"],
            ["Classification, full history", "Accuracy, macro F1, RMSE, MAE, R² on classes 1–5", "Persistence (accuracy 0.633 / 0.578 / 0.548). Learned models ~0.31–0.38.", "No"],
            ["Classification, from 4 April 2025", "Same five metrics on classes 1–5", "Only 72h Random Forest passed the gate (accuracy 0.568 vs 0.489). Walk-forward: 1 of 12 slots.", "No"],
        ],
    )
    caption(doc, "Table 33: Experiments I ran and what I kept.")
    body(
        doc,
        "I selected **continuous US EPA AQI regression on the post-4 April 2025 window**, with "
        "**LightGBM at 24 hours**, **Random Forest at 48 hours**, and a **top-3 tabular ensemble "
        "at 72 hours**. Those are the models behind aqi_forecaster_24h / 48h / 72h."
    )
    body(doc, "Why this, and not the others:")
    bullets(
        doc,
        [
            "A forecast that loses to “no change” should not be served. Full-history regression lost to persistence. Classification mostly did too.",
            "RMSE and MAE are AQI points, which is what a user feels. Post-break regression error is about **29–32**, roughly half of full-history persistence error (**67–92**).",
            "Low R² on the post-break window is expected on a calm summer test set. RMSE/MAE still improve versus persistence, and that is what I used to choose.",
            "The 4 April 2025 break is visible in the raw series. Training through it asks the model to learn a volatility that no longer exists.",
            "OpenWeather 1–5 classification would change the product (five buckets instead of a 0–500 index) and did not beat class-persistence in a stable way.",
            "The registered models are tabular or ensemble payloads that the dashboard can serve from one feature row.",
        ],
    )

    heading(doc, "23. Why Other Models Were Not Selected", 1)
    body(
        doc,
        "These reasons apply to the first notebook. The rule was the three-metric persistence gate."
    )
    heading(doc, "23.1 Ridge", 2)
    body(
        doc,
        "24/48/72 RMSE: 87.58 / 101.38 / 104.35 versus persistence 66.73 / 84.80 / 92.46. "
        "MAE and R² also fail at every horizon. Not registered."
    )
    heading(doc, "23.2 RandomForest", 2)
    body(
        doc,
        "24/48/72 RMSE: 84.91 / 118.05 / 116.12. All three metrics fail versus persistence at "
        "every horizon. Not registered."
    )
    heading(doc, "23.3 XGBoost", 2)
    body(
        doc,
        "24/48/72 RMSE: 74.36 / 97.21 / 111.99. Closest learned model at 24h and 48h, but RMSE "
        "and MAE remain higher than persistence and R² remains lower (24h R² 0.684 versus 0.745). "
        "Not registered."
    )
    heading(doc, "23.4 LightGBM", 2)
    body(
        doc,
        "24/48/72 RMSE: 76.46 / 101.16 / 112.21. Fails the three-metric gate. Not registered."
    )
    heading(doc, "23.5 Ensemble_top3_tabular", 2)
    body(
        doc,
        "24/48/72 RMSE: 77.09 / 98.11 / 107.62. Fails the three-metric gate. Not registered."
    )
    heading(doc, "23.6 LSTM", 2)
    body(
        doc,
        "24/48/72 RMSE: 232.03 / 234.14 / 235.97 with negative R². Not registered."
    )
    heading(doc, "23.7 GRU", 2)
    body(
        doc,
        "24/48/72 RMSE: 219.02 / 246.40 / 239.03 with negative R². Not registered."
    )
    heading(doc, "23.8 Prophet", 2)
    body(
        doc,
        "24/48/72 RMSE: 2167.32 / 2159.44 / 2150.78. Fails by a large margin. Not registered."
    )

    heading(doc, "24. Serving, Dashboard and Automation", 1)
    body(
        doc,
        "Streamlit reads Hopsworks directly. FastAPI is optional. The dashboard is a single page. "
        "I run it locally with Streamlit, or the full stack from one Docker image "
        "(dashboard on port 8501, API on port 8000). GitHub Actions "
        "refresh features hourly and retrain daily at 02:00 UTC. Feature-group version in those "
        "workflows is 4. The dashboard treats AQI **151** as the hazardous alert line."
    )
    body(
        doc,
        "The live app is **https://lahore-aqi-predictor.streamlit.app**."
    )

    heading(doc, "25. Key Findings", 1)
    bullets(
        doc,
        [
            "The raw snapshot has **46,419** rows and 16 stored columns from 27 November 2020 to 14 August 2026, with no duplicate timestamps and no missing values in stored columns.",
            "AQI is right-skewed (skewness 1.91), mean 247.8, median 177.",
            "pm2_5 correlation with AQI is **0.986**, which matches computing AQI from PM2.5.",
            "Smog-season (Oct–Jan) mean AQI is **399** versus **182** in other months.",
            "ADF on raw AQI rejects a unit root (statistic −11.76, p-value 1.13×10⁻²¹).",
            "The first notebook trained on aqi_features v3, 46,100 rows, 42 features, June-aligned split 27,947 / 8,075 / 10,078.",
            "Persistence test RMSE/MAE/R²: 24h **66.73 / 36.11 / 0.745**; 48h 84.80 / 45.78 / 0.588; 72h 92.46 / 51.52 / 0.510.",
            "**No learned model** beat persistence on RMSE, MAE and R² together. Nothing was registered from that run.",
            "Closest learned 24h model was XGBoost (RMSE 74.36).",
            "Prophet RMSE exceeded 2150 on every horizon. LSTM and GRU had negative R².",
            "OpenWeather breaks on 4 April 2025 (mean hourly |AQI change| ~46 → ~4.5). I now start training there.",
            "On that later window, selected RMSE was 28.88 / 30.54 / 31.65. Those figures are not from the first notebook.",
            "I also ran classification of OpenWeather’s 1–5 class on the same features. On full history, persistence accuracy was **0.633 / 0.578 / 0.548**; learned models sat around **0.31–0.38** and none passed the five-metric gate.",
            "On the post-break classification window, only **72h Random Forest** passed all five metrics (accuracy **0.568** versus **0.489**). Walk-forward held in **1 of 12** fold×horizon slots, so I did not adopt that path.",
            "The approach I selected is **continuous US EPA AQI regression from 4 April 2025**, with LightGBM (24h), Random Forest (48h) and a top-3 tabular ensemble (72h). Full-history regression and both classification windows were discarded because they did not beat persistence in a stable, serveable way.",
            "SHAP on the registered models: pressure (24h XGBoost), pm2_5 (48h RandomForest), month_cos / pm2_5 (72h ensemble members).",
        ],
    )

    heading(doc, "26. Limitations", 1)
    bullets(
        doc,
        [
            "AQI is PM2.5-only, not the official EPA maximum across six criteria pollutants.",
            "Feature group v3 (the first training table) was built on a gappy frame, so about 13% of lags/targets spanned the wrong number of hours.",
            "About 3,700 missing hours in 415 outages; I interpolate at most 6 hours and drop longer gaps.",
            "The first notebook registered nothing because no model beat persistence on all three metrics.",
            "Prophet, LSTM and GRU cannot be served from a one-row predict payload.",
            "The production-window test months are a relatively calm warm-season stretch; R² at 48h and 72h on that window is slightly negative even for the selected models.",
            "The EDA notebook in git has no saved plots; I regenerated those figures from the parquet.",
            "SHAP for the first notebook does not exist (nothing registered). Figures 12–21 use 400 boosting rows and 150 RandomForest rows on the post-break split.",
            "Classification metrics (Section 21) are on OpenWeather classes 1–5, not on 0–500 AQI, and are not mixed with the regression tables.",
            "The 72h Random Forest classification win on the post-break window did not hold in walk-forward (1 of 12 slots), so I did not register a classifier.",
        ],
    )

    heading(doc, "27. Conclusion", 1)
    body(
        doc,
        "I set out to forecast Lahore AQI at 24, 48 and 72 hours using OpenWeather pollution, "
        "Open-Meteo weather, Hopsworks storage, and a Streamlit dashboard with optional FastAPI. "
        "I compared two problem types — **continuous regression** of US EPA AQI from PM2.5 "
        "(0–500) and **classification** of OpenWeather’s 1–5 category — on both the full archive "
        "and the post-4 April 2025 window."
    )
    body(
        doc,
        "The raw series is right-skewed, tightly tied to PM2.5, worse in smog season, and strongly "
        "autocorrelated hour to hour. I engineered log pollutants, ACF-based lags, rolling windows, "
        "change rates, cyclical time features, a smog flag, weather interactions, and delta targets "
        "on an hourly grid."
    )
    body(
        doc,
        "On full-history regression (feature group v3, June-aligned split), **persistence had the "
        "best RMSE, MAE and R² at every horizon** (66.73 / 84.80 / 92.46). Closest learned 24h "
        "model was XGBoost (74.36). I registered nothing from that run."
    )
    body(
        doc,
        "On full-history classification, persistence accuracy was **0.633 / 0.578 / 0.548** and "
        "learned models sat around **0.31–0.38**. No classifier passed the five-metric gate. On "
        "the post-break classification window only **72h Random Forest** passed (accuracy 0.568 "
        "versus 0.489), and that win did not hold in walk-forward (1 of 12 slots)."
    )
    body(
        doc,
        "I selected **continuous US EPA AQI regression from 4 April 2025 onwards**: LightGBM at "
        "24 hours (RMSE **28.88**), Random Forest at 48 hours (**30.54**), and a top-3 tabular "
        "ensemble at 72 hours (**31.65**). I kept this path because those models beat persistence "
        "on AQI-point error, the regime break is visible in the raw series, 1–5 classification "
        "is a different product and was unstable, and the registered artefacts are serveable from "
        "one feature row. Full-history regression and both classification experiments remain in "
        "this report as discarded alternatives, not as competing production numbers."
    )
    body(
        doc,
        "SHAP on the models I actually serve is in Section 20.1. Pressure and 24-hour AQI change "
        "matter most for 24h XGBoost; PM2.5 for 48h RandomForest; the 72h ensemble members split "
        "between month_cos and PM2.5."
    )

    doc.save(OUT)
    print("wrote", OUT)


if __name__ == "__main__":
    build()
