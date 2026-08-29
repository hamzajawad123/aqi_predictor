"""Build reports/final_report.docx."""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

OUT = Path(__file__).resolve().parent / "final_report.docx"
NAVY = RGBColor(0x1B, 0x36, 0x5D)
ACCENT = RGBColor(0x2C, 0x5F, 0x8A)
MUTED = RGBColor(0x55, 0x55, 0x55)


def set_run(run, *, size=11, bold=False, italic=False, color=None, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "BFBFBF")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def para(doc, text, *, size=11, bold=False, italic=False, color=None, align="left", space_after=8, space_before=0, first_line=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.15
    if first_line and align == "left" and not bold:
        p.paragraph_format.first_line_indent = Cm(0.5)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, italic=italic, color=color)
    return p


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY if level == 1 else ACCENT
        run.font.name = "Calibri"
    h.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    h.paragraph_format.space_after = Pt(8)
    return h


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.first_line_indent = Cm(0)
        for run in p.runs:
            set_run(run, size=11)


def caption(doc, text):
    p = para(doc, text, size=9, italic=True, color=MUTED, first_line=False, space_after=12, space_before=2)
    p.paragraph_format.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run(run, size=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(hdr[i], "1B365D")
        set_cell_border(hdr[i])
    for r_i, row in enumerate(rows):
        cells = table.rows[r_i + 1].cells
        bg = "F4F7FA" if r_i % 2 == 0 else "FFFFFF"
        for c_i, val in enumerate(row):
            cells[c_i].text = ""
            p = cells[c_i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_run(run, size=9)
            shade_cell(cells[c_i], bg)
            set_cell_border(cells[c_i])
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def body(doc, text):
    p = para(doc, text, align="justify", first_line=True)
    p.paragraph_format.first_line_indent = Cm(0.5)
    return p


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run("Pearls AQI Predictor  ·  10Pearls Data Science Internship  ·  Confidential student report")
        set_run(run, size=8, color=MUTED, italic=True)

    # Cover
    for _ in range(3):
        doc.add_paragraph()
    t = para(doc, "10PEARLS  ·  DATA SCIENCE INTERNSHIP", size=12, bold=True, color=ACCENT, align="center", first_line=False, space_after=6)
    t.paragraph_format.first_line_indent = Cm(0)
    para(doc, "TECHNICAL PROJECT REPORT", size=14, bold=True, color=NAVY, align="center", first_line=False, space_after=18)
    para(
        doc,
        "Pearls AQI Predictor:\nA Serverless Feature–Training–Inference Pipeline\nfor Three-Day-Ahead Air Quality Forecasting in Lahore",
        size=20,
        bold=True,
        color=NAVY,
        align="center",
        first_line=False,
        space_after=24,
    )
    para(
        doc,
        "Continuous US EPA AQI (0–500) from PM2.5  ·  24 / 48 / 72 hour horizons",
        size=12,
        italic=True,
        color=MUTED,
        align="center",
        first_line=False,
        space_after=36,
    )
    meta = [
        "Author: ________________________________",
        "Programme: 10Pearls Data Science Internship",
        "Repository: https://github.com/hamzajawad123/aqi_predictor",
        "Document date: 27 August 2026",
        "Status: Core methodology complete. Live deployment URL to be inserted after hosting.",
    ]
    for line in meta:
        p = para(doc, line, size=12, align="center", first_line=False, space_after=4)
        p.paragraph_format.first_line_indent = Cm(0)

    doc.add_page_break()

    heading(doc, "Abstract", 1)
    body(
        doc,
        "This report documents the end-to-end design and evaluation of a three-day-ahead "
        "Air Quality Index (AQI) forecast for Lahore, built as a serverless Feature / Training / "
        "Inference (FTI) pipeline. Pollution concentrations are obtained exclusively from the "
        "OpenWeather Air Pollution API. Meteorological covariates are obtained exclusively from "
        "Open-Meteo. Features and trained models are stored in Hopsworks. Inference is served "
        "by FastAPI and consumed by a Streamlit dashboard. GitHub Actions refresh features hourly "
        "and can retrain daily."
    )
    body(
        doc,
        "The prediction target is a continuous US EPA AQI in the range 0–500, computed from "
        "PM2.5 using the 2024 breakpoint table. OpenWeather’s native main.aqi field is a coarse "
        "1–5 category and is retained only as a reference column; it is not the production target. "
        "Tabular and recurrent models are trained on AQI deltas and scored on reconstructed "
        "absolute AQI so that every family is comparable to a persistence baseline "
        "(“AQI in N hours equals AQI now”). A candidate is registered in Hopsworks only if it "
        "beats persistence on RMSE, MAE and R² simultaneously."
    )
    body(
        doc,
        "Two training windows were compared. Approach 1 uses only data after a documented "
        "regime break on 4 April 2025. On that window the selected models beat persistence, "
        "with day-1 / day-2 / day-3 RMSE of 28.88 / 30.54 / 31.65 AQI points. R² is near zero "
        "because the held-out months are a calm summer period with low AQI variance; that does "
        "not invalidate the error reduction versus baseline. Approach 2 uses the full archive "
        "from late 2020. R² looks stronger (approximately 0.75 / 0.59 / 0.51), but persistence "
        "wins every horizon and no model is eligible for registration. Following mentor review, "
        "Approach 1 is the production choice. A parallel 1–5 classification experiment was also "
        "run and was not adopted."
    )

    heading(doc, "Contents", 1)
    toc_items = [
        "1. Introduction and problem framing",
        "2. Objectives, scope and non-goals",
        "3. System architecture",
        "4. Data sources and API selection",
        "5. Target definition: why continuous EPA AQI",
        "6. Collection, alignment, cleaning and validation",
        "7. Exploratory data analysis",
        "8. Feature engineering",
        "9. Experimental protocol",
        "10. Models evaluated",
        "11. Results: Approach 1 versus Approach 2",
        "12. Classification experiment (not adopted)",
        "13. Final model selection and rationale",
        "14. Serving, dashboard and automation",
        "15. Blockers encountered and how they were resolved",
        "16. Limitations and remaining work",
        "17. Conclusion",
        "Appendix A. Classification metric tables",
        "Appendix B. Repository map",
    ]
    for item in toc_items:
        p = para(doc, item, first_line=False, space_after=2)
        p.paragraph_format.first_line_indent = Cm(0)

    heading(doc, "1. Introduction and problem framing", 1)
    body(
        doc,
        "Lahore experiences severe particulate pollution, especially in the October–January "
        "smog season. A three-day-ahead AQI forecast is operationally useful: it supports "
        "outdoor-activity decisions, school and workplace planning, and simple hazardous-air "
        "alerts. The internship brief asked for a serverless FTI pipeline that predicts AQI "
        "and reports RMSE, MAE and R², together with exploratory analysis, feature engineering, "
        "model comparison, an inference API, a dashboard, and scheduled pipelines."
    )
    body(
        doc,
        "The scientific difficulty is not “fit a model to a large table.” Hourly AQI is strongly "
        "autocorrelated. A naïve persistence forecast is therefore a serious baseline, not a "
        "formality. Any learned model that cannot beat “tomorrow looks like today” on error "
        "metrics is not worth serving, even if its R² looks attractive on a high-variance window. "
        "That principle governs every registration decision in this project."
    )
    body(
        doc,
        "Work proceeded in this order: lock one pollution source and one weather source; define "
        "a continuous AQI target; persist a reproducible raw snapshot; run EDA before feature "
        "engineering; engineer lags, rolling statistics, time encodings and delta targets on a "
        "strict hourly grid; compare model families against persistence on two data windows; "
        "wire serving and CI; and, after mentor feedback, freeze Approach 1 as production."
    )

    heading(doc, "2. Objectives, scope and non-goals", 1)
    heading(doc, "2.1 Objectives", 2)
    bullets(
        doc,
        [
            "Forecast Lahore AQI at 24, 48 and 72 hours (day 1, day 2 and day 3).",
            "Keep pollution and weather each on a single source in both historical backfill and hourly refresh, to avoid train/serving skew.",
            "Store engineered features in Hopsworks and register only models that beat persistence on RMSE, MAE and R².",
            "Serve forecasts through FastAPI and present them in Streamlit, including a hazardous-AQI alert.",
            "Automate hourly feature updates and optional daily retraining with GitHub Actions.",
        ],
    )
    heading(doc, "2.2 Scope boundaries", 2)
    bullets(
        doc,
        [
            "City: Lahore (31.5497° N, 74.3436° E). The pipeline is location-parameterised but was not multi-city trained.",
            "AQI is computed from PM2.5 only (2024 US EPA breakpoints). A full multi-pollutant EPA max was not implemented.",
            "OpenWeather’s own 4-day pollution forecast is fetched in code for possible later comparison; it is not the training target.",
            "Prophet, LSTM and GRU are trained and scored, but /predict serves only tabular or ensemble payloads that consume one feature row.",
        ],
    )

    heading(doc, "3. System architecture", 1)
    body(
        doc,
        "The system follows the FTI pattern required by the internship. The feature pipeline "
        "fetches, validates, engineers and writes. The training pipeline reads the feature group, "
        "fits models, applies a persistence gate, and may write to the model registry. The "
        "inference pipeline loads the latest feature row and the registered per-horizon models "
        "and returns three point forecasts plus a hazardous flag (AQI ≥ 151, Unhealthy for "
        "Sensitive Groups and worse on the continuous EPA scale)."
    )
    para(doc, "Data flow (production path)", bold=True, first_line=False, space_after=6)
    para(
        doc,
        "OpenWeather (pollution)  +  Open-Meteo (weather)\n"
        "        →  merge on UTC hour  →  validate  →  raw parquet snapshot\n"
        "        →  feature engineering  →  Hopsworks feature group aqi_features v4\n"
        "        →  training pipeline  →  Hopsworks model registry  aqi_forecaster_{24,48,72}h\n"
        "        →  FastAPI  GET /predict  →  Streamlit (EDA, forecast, alerts, model performance)",
        size=10,
        italic=True,
        first_line=False,
        space_after=12,
    )
    body(
        doc,
        "Local Docker Compose runs the API and dashboard from one image (ports 8000 and 8501). "
        "Streamlit reads Hopsworks directly; FastAPI is an optional HTTP wrapper around the same "
        "serving code. The public dashboard is on Streamlit Community Cloud. Hopsworks remains "
        "the feature store and registry in every environment."
    )

    heading(doc, "4. Data sources and API selection", 1)
    heading(doc, "4.1 Pollution: OpenWeather Air Pollution API", 2)
    body(
        doc,
        "I selected OpenWeather as the sole pollution source for current, forecast and historical "
        "endpoints. Historical air-pollution access on the free tier begins on 27 November 2020, "
        "which also became DATA_START_DATE so that weather and pollution share the same first day. "
        "The payload includes pollutant concentrations (CO, NO, NO2, O3, SO2, PM2.5, PM10, NH3) "
        "and a field named main.aqi. That field is documented as an index from 1 (Good) to 5 "
        "(Very Poor). It is not a continuous AQI and must not be confused with the 0–500 scale "
        "shown on public dashboards such as AQICN."
    )
    body(
        doc,
        "Using one pollution vendor everywhere matters more than vendor branding. Mixing, for "
        "example, OpenWeather history with a different live feed would silently change the "
        "meaning of PM2.5 between training and serving. The feature store cannot fix that class "
        "of bug; source discipline can."
    )
    heading(doc, "4.2 Weather: Open-Meteo", 2)
    body(
        doc,
        "I selected Open-Meteo for temperature, relative humidity, wind speed, wind direction "
        "and surface pressure. It requires no API key, provides a consistent archive and recent "
        "forecast/hourly API, and is sufficient for the meteorological covariates used in this "
        "project. An earlier draft used OpenWeather for live weather and Open-Meteo for backfill. "
        "That split is exactly train/serving skew: the same column name would not be the same "
        "measurement process. Weather is therefore Open-Meteo in both the hourly path and the "
        "historical path."
    )
    heading(doc, "4.3 Sources that were considered and not used as the target", 2)
    bullets(
        doc,
        [
            "OpenWeather main.aqi (1–5): stored as openweather_aqi_category; never the regression target.",
            "AQICN screenshot values in the brief: used only to confirm that “AQI” means a continuous index, not to scrape a second pollution feed.",
            "OpenWeather weather endpoints: not used in production, to keep one weather source.",
        ],
    )

    heading(doc, "5. Target definition: why continuous EPA AQI", 1)
    body(
        doc,
        "The brief asks to predict AQI and to report RMSE, MAE and R². Those metrics assume a "
        "numeric target with meaningful distances. A 1–5 category treated as a number produces "
        "artificially small errors and inflated R², because the label has only five values. "
        "Colleagues who model main.aqi are solving a different, easier problem. I compute AQI "
        "from PM2.5 with the US EPA 2024 revised breakpoints (Federal Register reconsideration "
        "of the PM NAAQS): AQI 50 maps to 9.0 µg/m³, 100 and 150 remain at 35.4 and 55.4 µg/m³, "
        "and 200 / 300 / 500 map to 125.4 / 225.4 / 325.4 µg/m³."
    )
    body(
        doc,
        "The official EPA method truncates concentration to one decimal place before breakpoint "
        "lookup, leaving intentional 0.1 µg/m³ gaps between bands (9.0 | 9.1, 35.4 | 35.5, …). "
        "Skipping truncation caused a real bug: a raw float such as 9.0989 fell in a gap and "
        "produced a nonsensical negative AQI. The implementation now truncates first. Unit tests "
        "cover EPA worked examples, monotonicity on a dense grid, and the specific gap value "
        "that previously failed. Concentrations above the top breakpoint are extrapolated rather "
        "than hard-capped at 500, so extreme Lahore smog days remain ordered instead of collapsing "
        "to a single ceiling."
    )
    body(
        doc,
        "Simplification, stated explicitly: a complete EPA AQI is the maximum across six criteria "
        "pollutants. This project uses PM2.5 only, because PM2.5 dominates Lahore’s smog regime "
        "and correlates extremely strongly with the constructed AQI series. Extending to a true "
        "multi-pollutant max would require verified breakpoint tables for the other species."
    )

    heading(doc, "6. Collection, alignment, cleaning and validation", 1)
    heading(doc, "6.1 Time alignment", 2)
    body(
        doc,
        "Both APIs are normalised to UTC before the merge. The join key is the hourly timestamp. "
        "No manual timezone offset is applied at merge time. A separate bug existed when converting "
        "naive timestamps to UNIX seconds for OpenWeather’s history endpoint: pandas Timestamp.timestamp() "
        "treats a naive value as local time, which shifted the requested window by the host UTC offset "
        "(UTC+5 on a Pakistan machine). Naive hours in this project are UTC wall-clock hours, so they "
        "are localised to UTC before conversion. Tests lock that behaviour."
    )
    heading(doc, "6.2 Raw snapshot", 2)
    body(
        doc,
        "Every hourly and backfill run upserts a local parquet file, data/raw/aqi_raw_merged.parquet. "
        "EDA and later feature experiments read this snapshot, not the engineered Hopsworks table. "
        "That separation prevents “EDA on leaked targets” and keeps a reproducible pre-FE archive "
        "when the feature-group schema changes (v1/v2 rollback, v3 gappy-frame features, v4 hourly grid)."
    )
    heading(doc, "6.3 Validation before the feature store", 2)
    body(
        doc,
        "validate_raw_data runs before insert. It checks required columns, drops duplicate timestamps, "
        "drops rows with nulls in required fields, and drops out-of-range sensor/API glitches using "
        "wide physical bounds (for example AQI 0–1000, PM2.5 0–2000 µg/m³, temperature −30–60 °C). "
        "Bounds are intentionally wide so genuine smog extremes are not treated as errors. One bad "
        "hour drops that row; it does not abort the entire batch, unless raise_on_error is set for "
        "a hard CI fail."
    )
    heading(doc, "6.4 Missing hours and the hourly grid", 2)
    body(
        doc,
        "The raw history contains on the order of 3,700 missing hours across hundreds of outages. "
        "Lags, rolling windows and targets were originally implemented as row shifts. On a gappy "
        "frame, shift(−24) spanned 24 hours for only about 87% of rows and up to 264 hours for the "
        "rest — those rows were trained against the wrong future. The fix is to_hourly_grid: reindex "
        "onto a strict hourly calendar first. Outages of at most six hours, end to end, are linearly "
        "interpolated. Longer gaps stay NaN and are dropped. Pandas’ own interpolate(limit=6) would "
        "have filled the first six hours of a 20-hour hole toward a reading 20 hours away; that was "
        "rejected as inventing a trajectory. Integer-like columns are rounded immediately after "
        "interpolation so Hopsworks bigint schema and derived lags stay consistent."
    )

    heading(doc, "7. Exploratory data analysis", 1)
    body(
        doc,
        "EDA lives in notebooks/01_eda.ipynb and is required before feature-engineering changes. "
        "It covers structure checks, univariate distributions, bivariate relationships, a smog-versus-"
        "normal seasonal split (October–January), seasonal decomposition, ADF stationarity, and "
        "ACF/PACF. A closing “Findings for FE” cell prints the quantities that actually entered "
        "the feature pipeline so the notebook cannot drift from the code without being obvious."
    )
    heading(doc, "7.1 Findings that changed the feature set", 2)
    bullets(
        doc,
        [
            "Pollutant concentrations are strongly right-skewed (skew ≫ 1). Production features use log1p on CO, NO, NO2, O3, SO2, PM2.5, PM10 and NH3.",
            "ACF/PACF on raw AQI justified lags at 1, 3, 6, 24 and 168 hours (day and week), not an arbitrary dense lag set.",
            "Rolling windows of 3, 6 and 24 hours capture short-term level and volatility.",
            "AQI exhibits a multi-year downward shift in volatility and level. Absolute future AQI as a train target fights that drift; delta targets (future AQI minus current AQI) became the primary supervised target for tabular and recurrent models.",
            "Lahore smog season (Oct–Jan) is a first-class flag and the stratification key for smog-versus-normal evaluation, using one shared month tuple so the two cannot silently disagree.",
        ],
    )
    heading(doc, "7.2 The 4 April 2025 regime break", 2)
    body(
        doc,
        "OpenWeather’s air-pollution archive changes character on 4 April 2025. Mean hourly "
        "|AQI change| collapses from roughly 46 to roughly 4.5 and does not recover; the same "
        "break appears in PM2.5. Data before that date are a different generating process. "
        "Training across the break teaches models a volatility that no longer exists. Empirically, "
        "that is why every learned model lost to persistence on a 2025–2026 test period when the "
        "full history was used. REGIME_BREAK_DATE and the default TRAIN_START_DATE are therefore "
        "set to 2025-04-04. Empty TRAIN_START_DATE remains available as an explicit “train on all "
        "history” experiment — Approach 2 in this report."
    )

    heading(doc, "8. Feature engineering", 1)
    body(
        doc,
        "build_feature_set is the single function used conceptually for both writing the feature "
        "group and understanding the training matrix, so serving cannot drift from training. "
        "Training mode drops any row with a NaN, which removes the start of the series (incomplete "
        "lags) and the last 72 hours (unknown targets). Inference mode drops only incomplete "
        "feature columns and keeps rows whose 72-hour target is still in the future — those are "
        "exactly the rows /predict must score. Target columns exist in both modes (NaN when the "
        "future is unknown) so the Hopsworks schema does not fork by caller."
    )
    heading(doc, "8.1 Feature groups", 2)
    bullets(
        doc,
        [
            "Time: hour, day of week, month, weekend flag; cyclical hour_sin/cos and month_sin/cos so 23:00 is close to 00:00.",
            "Season: is_smog_season for October–January.",
            "Transform: log1p pollutants.",
            "Lags of AQI: 1, 3, 6, 24, 168 hours.",
            "Rolling mean/std/min/max of AQI: 3, 6, 24 hours.",
            "Change rates: 1-hour and 24-hour AQI differences (explicitly required by the brief).",
            "Weather interactions: wind_speed × PM2.5 and humidity × PM2.5 (dispersion and secondary particle formation).",
            "Targets: aqi_target_{24,48,72}h and aqi_delta_{24,48,72}h.",
            "Reference only: openweather_aqi_category is stored but dropped from model inputs.",
        ],
    )
    heading(doc, "8.2 Why deltas, then score on absolute AQI", 2)
    body(
        doc,
        "Models see aqi_delta_h as y. At evaluation and serving time the prediction is reconstructed "
        "as current AQI + shrinkage × predicted delta. Persistence is exactly shrinkage = 0. "
        "Prophet is the exception: it is fit on absolute AQI as a level series. All families are "
        "scored on absolute future AQI so RMSE/MAE/R² remain in AQI points and remain comparable."
    )
    heading(doc, "8.3 Delta shrinkage", 2)
    body(
        doc,
        "Predicted deltas are noisy. On hours when AQI barely moves — the majority of the recent "
        "regime — an over-confident delta is worse than predicting zero change. A shrinkage factor "
        "λ ∈ {0.00, 0.05, …, 1.00} is fitted on validation only (never on test). λ = 0 reproduces "
        "persistence; λ = 1 is the raw model. The search therefore cannot select a validation "
        "fit worse than persistence. Serving stores and applies the same λ so production metrics "
        "match the registered numbers."
    )
    heading(doc, "8.4 Feature-group versions", 2)
    body(
        doc,
        "v1/v2 are prior engineered tables kept for rollback. v3 has delta targets but was built "
        "on the gappy frame, so about 13% of lags/targets spanned the wrong number of hours. "
        "v4 is the production group: same feature definitions on the strict hourly grid. "
        "FEATURE_GROUP_VERSION defaults to 4 and is pinned to 4 in GitHub Actions."
    )

    heading(doc, "9. Experimental protocol", 1)
    heading(doc, "9.1 Splits", 2)
    body(
        doc,
        "The default split is chronological and season-aligned: validation and test windows snap "
        "to 1 June so a smog season is not cut in half. If the selected training window is too "
        "short for a full June-to-June layout (Approach 1), the code falls back to a 70 / 15 / 15 "
        "chronological split and logs that fact. Optuna uses TimeSeriesSplit on the training fold "
        "only. Walk-forward checks were used as a robustness diagnostic, not as the registration rule."
    )
    heading(doc, "9.2 Persistence gate (registration rule)", 2)
    body(
        doc,
        "For each horizon independently, every model is compared with persistence on RMSE, MAE "
        "and R². All three must improve. Among models that pass, the lowest RMSE wins (MAE then "
        "R² as tie-breakers). The winner is stored as aqi_forecaster_{h}h. If the statistical "
        "winner is Prophet, LSTM or GRU — families that cannot consume the single flat row that "
        "/predict provides — the pipeline falls back to the best serveable (tabular or ensemble) "
        "candidate that still beats persistence, and logs the substitution. If nobody passes, "
        "nothing is registered for that horizon."
    )
    heading(doc, "9.3 Metrics interpretation", 2)
    body(
        doc,
        "RMSE and MAE are in AQI points and are the operational scores. R² = 1 − MSE / Var(y_test). "
        "On a calm test window, Var(y_test) is small, so a respectable RMSE still yields R² near "
        "zero. Chasing R² of 0.8 on that window would require an RMSE far below anything observed. "
        "Conversely, a high-variance window makes persistence look statistically strong (high R²) "
        "while leaving large AQI-point errors. Mentor feedback confirmed this reading: low R² in "
        "a calm window is acceptable when RMSE/MAE are good and the model beats persistence."
    )

    heading(doc, "10. Models evaluated", 1)
    body(
        doc,
        "The production training path (src/training_pipeline.py and notebooks/02_training.ipynb) "
        "evaluates the following families on each horizon."
    )
    add_table(
        doc,
        ["Family", "Train target", "Notes"],
        [
            ["Persistence", "δ = 0", "Baseline; equivalent to “future AQI = current AQI”"],
            ["Prophet", "Absolute AQI", "Level series; not serveable via one-row /predict"],
            ["Ridge", "Delta", "Linear; Optuna-tuned"],
            ["Random Forest", "Delta", "Optuna-tuned"],
            ["XGBoost", "Delta", "Optuna-tuned"],
            ["LightGBM", "Delta", "Optuna-tuned; L1 variant among candidates"],
            ["Mean ensembles", "Absolute reconstructions", "Simple averages of tabular predictions"],
            ["LSTM / GRU", "Delta, 24h windows", "Scored; not the default /predict payload"],
            ["Shrinkage variants", "Scaled delta", "λ fit on validation for each delta model"],
        ],
    )
    caption(doc, "Table 1. Model families in the regression path.")
    body(
        doc,
        "Hyperparameters for Ridge, Random Forest, XGBoost and LightGBM are chosen with Optuna "
        "on time-series cross-validation. Recurrent nets use a 24-hour sequence length matching "
        "src/utils/sequences.py. SHAP summary plots are generated for a winning tree model per "
        "horizon when applicable (reports/shap_summary_{24,48,72}h.png)."
    )

    heading(doc, "11. Results: Approach 1 versus Approach 2", 1)
    body(
        doc,
        "This comparison is the decision the mentor asked to record in the report. Both approaches "
        "use the same feature definitions, the same model families, the same persistence gate, and "
        "the same 24 / 48 / 72 hour horizons. They differ only in how much history is shown to "
        "the models."
    )

    heading(doc, "11.1 Approach 1 — current regime (from 4 April 2025)", 2)
    body(
        doc,
        "Training starts at the regime break. The test window falls in a relatively calm "
        "warm-season period (low hour-to-hour AQI movement). Learned models beat persistence "
        "on all three metrics. Selected winners:"
    )
    add_table(
        doc,
        ["Horizon", "Selected model", "RMSE", "MAE", "R²"],
        [
            ["24 h (day 1)", "LightGBM (L1)", "28.88", "22.63", "0.064"],
            ["48 h (day 2)", "Random Forest", "30.54", "24.89", "−0.055"],
            ["72 h (day 3)", "Ensemble (top-3 tabular)", "31.65", "26.07", "−0.098"],
        ],
    )
    caption(doc, "Table 2. Approach 1 selected models (absolute AQI points).")
    add_table(
        doc,
        ["Horizon", "Persistence RMSE", "Selected-model RMSE"],
        [
            ["24 h", "34.84", "28.88"],
            ["48 h", "40.94", "30.54"],
            ["72 h", "40.30", "31.65"],
        ],
    )
    caption(
        doc,
        "Table 3. Approach 1 persistence versus selected-model RMSE on the same test window. Persistence is worse at every horizon.",
    )
    body(
        doc,
        "Interpretation: typical error is about 23–26 AQI points (MAE). That is a usable "
        "day-ahead residual on a 0–500 scale. R² near zero (slightly negative at 48 h and 72 h) "
        "means the model does not explain residual variance beyond the mean in a low-variance "
        "test set; it still reduces RMSE and MAE relative to persistence, which is the operational "
        "claim. These three models are the ones designated for Hopsworks as aqi_forecaster_24h, "
        "aqi_forecaster_48h and aqi_forecaster_72h."
    )

    heading(doc, "11.2 Approach 2 — full history (from late 2020)", 2)
    body(
        doc,
        "The same pipeline was run with the full archive (notebooks/02_training.ipynb as executed "
        "on Colab against an earlier feature-group version). Persistence won every horizon. "
        "No model beat persistence on all three metrics. Nothing was registered."
    )
    add_table(
        doc,
        ["Horizon", "Winner", "RMSE", "MAE", "R²"],
        [
            ["24 h", "Persistence", "66.73", "36.11", "0.745"],
            ["48 h", "Persistence", "84.80", "45.78", "0.588"],
            ["72 h", "Persistence", "92.46", "51.52", "0.510"],
        ],
    )
    caption(doc, "Table 4. Approach 2: persistence is the winner. Learned models were worse (higher error) despite lower R² looking “harder.”")
    body(
        doc,
        "R² of 0.75 / 0.59 / 0.51 looks closer to naïve academic targets of 0.8 / 0.6 / 0.5. "
        "Those numbers are easy when y_test has large variance (older volatile winters mixed into "
        "the series) and the predictor is essentially “today’s AQI.” RMSE of 67–92 AQI points is "
        "a worse forecast than Approach 1’s 29–32. Best gradient-boosted and linear models in "
        "that run still lost to persistence. I therefore do not register Approach 2 models."
    )

    heading(doc, "11.3 Head-to-head", 2)
    add_table(
        doc,
        ["Criterion", "Approach 1 (regime)", "Approach 2 (full history)"],
        [
            ["Training start", "4 April 2025", "November/December 2020"],
            ["Beats persistence?", "Yes", "No"],
            ["Day-1 RMSE (best)", "28.88 (LightGBM)", "66.73 (persistence)"],
            ["Day-1 R²", "0.06 (calm test)", "0.75 (persistence)"],
            ["Register?", "Yes", "No"],
            ["Matches current data-generating process?", "Yes", "No — pre-break volatility"],
        ],
    )
    caption(doc, "Table 5. Decision matrix used for the final choice.")

    heading(doc, "12. Classification experiment (not adopted)", 1)
    body(
        doc,
        "After colleague suggestions to treat OpenWeather’s 1–5 index as the target, I ran a "
        "parallel classification study without changing the production regression pipeline. "
        "The target was future openweather_aqi_category at 24 / 48 / 72 hours. Features were "
        "the same production matrix. Models included persistence (today’s class), ordinal "
        "logistic regression, Random Forest / XGBoost / LightGBM classifiers, a majority-vote "
        "ensemble, LSTM and GRU with a 5-way softmax. A model “won” only if it beat persistence "
        "on accuracy, macro F1, RMSE, MAE and R² together (RMSE/MAE/R² computed on the 1–5 labels). "
        "Nothing was written to the registry."
    )
    body(
        doc,
        "Full history (December 2020–August 2026, ~39,060 hours): persistence accuracy was "
        "0.633 / 0.578 / 0.548 at 24 / 48 / 72 hours. Learned models clustered around 0.31–0.38. "
        "No model passed the five-metric gate."
    )
    body(
        doc,
        "Regime window (April 2025–August 2026, ~9,866 hours): results were closer. Only "
        "72-hour Random Forest passed all five metrics (accuracy 0.568 versus persistence 0.489; "
        "RMSE 0.746 versus 0.895). Forty-eight-hour GRU looked strong on accuracy (0.568 versus "
        "0.486) but lost on macro F1 (0.292 versus 0.298). Walk-forward (four folds, Random Forest "
        "versus persistence) passed the five-metric test in 1 of 12 fold×horizon slots. The 72-hour "
        "win is therefore not stable."
    )
    body(
        doc,
        "Conclusion: classifying 1–5 is a different product and does not beat “today’s class” "
        "in a reliable way. I did not replace the 0–500 regression path. Full classification "
        "RMSE/MAE/R² tables are in Appendix A."
    )

    heading(doc, "13. Final model selection and rationale", 1)
    body(
        doc,
        "I select Approach 1 as the production system: LightGBM at 24 hours, Random Forest at "
        "48 hours, and a top-3 tabular ensemble at 72 hours, each trained on post-4-April-2025 "
        "data with delta targets, validation-fitted shrinkage, and a persistence gate."
    )
    para(doc, "Reasons, in order of importance", bold=True, first_line=False, space_after=6)
    bullets(
        doc,
        [
            "The models beat persistence on RMSE, MAE and R². Approach 2’s models do not. A forecast that loses to “no change” should not be served.",
            "RMSE/MAE are the quantities a user feels (AQI points). Approach 1 errors are roughly half of Approach 2’s persistence errors.",
            "Low R² on Approach 1 is expected on a calm summer test set and was explicitly accepted by the mentor: “Low R² in a calm window is fine; RMSE/MAE are good and that's what matters.”",
            "Approach 2’s high R² is persistence explaining high historical variance, not a better learned forecast.",
            "The 4 April 2025 break is visible in the raw series. Training through it asks the model to unlearn a process that no longer exists.",
            "Classification of OpenWeather 1–5 did not produce a stable improvement over class-persistence and would change the product definition.",
            "Serveability: registered artefacts are tabular or ensemble payloads compatible with GET /predict.",
        ],
    )
    body(
        doc,
        "I do not claim a high-R² 72-hour AQI model on calm data. I claim a pipeline that is "
        "honest about baselines, refuses to register models that lose to persistence, and ships "
        "the horizon-wise winners that actually reduce error on the current regime."
    )

    heading(doc, "14. Serving, dashboard and automation", 1)
    heading(doc, "14.1 FastAPI", 2)
    body(
        doc,
        "api/main.py loads aqi_forecaster_{24,48,72}h from the registry (cached) and the latest "
        "feature-store row. Delta models reconstruct absolute AQI with the stored shrinkage. "
        "The JSON payload includes city, current AQI, three horizon forecasts, model names, and "
        "hazardous_alert. If a horizon has no registered model, that slot returns an explicit "
        "error string rather than a silent zero."
    )
    heading(doc, "14.2 Streamlit", 2)
    body(
        doc,
        "The dashboard is Streamlit, not a separate React/Vercel frontend. Pages: Home, EDA "
        "(local raw parquet), Forecast, Alerts, and Model Performance (comparison CSVs and SHAP "
        "images under reports/). Forecast and alerts call the API. This matches the internship "
        "stack and avoids rewriting the UI solely for hosting fashion."
    )
    heading(doc, "14.3 GitHub Actions", 2)
    body(
        doc,
        "feature_pipeline.yml runs hourly (and on workflow_dispatch): python -m src.feature_pipeline. "
        "training_pipeline.yml is scheduled daily at 02:00 UTC. Required secrets are "
        "OPENWEATHER_API_KEY, HOPSWORKS_API_KEY and HOPSWORKS_PROJECT_NAME. Feature-group version "
        "is pinned to 4 in the workflow files so an old repository secret cannot retarget inserts "
        "at a pre-delta schema."
    )
    heading(doc, "14.4 Tests", 2)
    body(
        doc,
        "pytest covers EPA AQI calculation (including the breakpoint-gap regression), data "
        "validation, evaluation helpers, feature engineering, and feature-pipeline time handling. "
        "Tests are the safety net for the bugs in Section 15."
    )

    heading(doc, "15. Blockers encountered and how they were resolved", 1)

    heading(doc, "15.1 Wrong target: treating main.aqi as AQI", 2)
    body(
        doc,
        "Blocker: OpenWeather names a 1–5 field “aqi”. Training regression against it would have "
        "matched neither the brief’s 0–500 intuition nor RMSE/MAE as AQI-point errors. "
        "Fix: implement US EPA PM2.5 AQI; store the 1–5 field separately; never use it as the "
        "regression label. Later, run an explicit classification experiment and reject it on evidence."
    )
    heading(doc, "15.2 EPA breakpoint gaps producing negative AQI", 2)
    body(
        doc,
        "Blocker: values such as PM2.5 = 9.0989 fell between 9.0 and 9.1 and escaped every band. "
        "Fix: truncate to one decimal per EPA method; add grid tests for non-negativity and monotonicity."
    )
    heading(doc, "15.3 Train/serving skew on weather", 2)
    body(
        doc,
        "Blocker: live weather from OpenWeather and historical weather from Open-Meteo. "
        "Fix: Open-Meteo for weather in both paths; OpenWeather for pollution only."
    )
    heading(doc, "15.4 Naive timestamps converted as local time", 2)
    body(
        doc,
        "Blocker: OpenWeather history windows were shifted by the host timezone. "
        "Fix: _unix_utc localises naive timestamps to UTC before .timestamp(); unit tests included."
    )
    heading(doc, "15.5 Row shifts on a gappy hourly series", 2)
    body(
        doc,
        "Blocker: ~13% of “24-hour” targets were not 24 hours. Worst cases spanned many days. "
        "Fix: strict hourly reindex; interpolate only ≤6-hour complete outages; drop the rest; "
        "promote feature group to v4."
    )
    heading(doc, "15.6 Hourly CI inserted zero rows", 2)
    body(
        doc,
        "Blocker: lookback was 200 hours. A row needs 168 hours of lag history and 72 hours of "
        "future for training-style drops; 200 hours cannot satisfy both, so build_feature_set "
        "returned an empty frame and inserts did nothing. Fix: lookback = max(lags) + max(horizons) "
        "+ margin (336 hours / 14 days). Inference mode already keeps rows with unknown targets; "
        "the longer window supplies lag history for the latest hour."
    )
    heading(doc, "15.7 Hopsworks schema mismatch in GitHub Actions", 2)
    body(
        doc,
        "Blocker: hourly Actions failed on insert of aqi_delta_* into an old feature-group version "
        "(a FEATURE_GROUP_VERSION secret still pointed at a pre-delta group). Hopsworks does not "
        "add columns on insert. Fix: pin FEATURE_GROUP_VERSION=4 in both workflows; fail fast with "
        "an explicit message if the opened group lacks delta columns; keep older versions as rollback."
    )
    heading(doc, "15.8 Hopsworks host instability", 2)
    body(
        doc,
        "Blocker: without an explicit host the SDK sometimes resolved a hostname that did not "
        "match the project region (for example c.app.hopsworks.ai versus eu-west.cloud.hopsworks.ai). "
        "Fix: pin HOPSWORKS_HOST in config (default eu-west.cloud.hopsworks.ai)."
    )
    heading(doc, "15.9 Integer interpolation versus bigint schema", 2)
    body(
        doc,
        "Blocker: interpolating across a gap made originally integer columns fractional; Hopsworks "
        "typed them as bigint. Fix: round immediately after the grid step; recast to int64 after NaNs "
        "are dropped; cast known integer feature columns before insert."
    )
    heading(doc, "15.10 Full-history models losing to persistence", 2)
    body(
        doc,
        "Blocker: a long archive looked “more data is better” but produced models worse than "
        "persistence and misleadingly high R². Fix: measure the regime break; default TRAIN_START_DATE "
        "to 4 April 2025; keep the full-history run as Approach 2 in this report rather than as production."
    )
    heading(doc, "15.11 R² pressure versus a calm test window", 2)
    body(
        doc,
        "Blocker: stakeholder expectation of R² ≈ 0.8 / 0.6 / 0.5 on 24 / 48 / 72 hours. On the "
        "Approach 1 test months that target is not identified with a good forecast. Fix: report "
        "RMSE/MAE as primary; explain R² algebraically; confirm with mentor that beating persistence "
        "on a calm window is the correct bar."
    )
    heading(doc, "15.12 Models that win statistically but cannot be served", 2)
    body(
        doc,
        "Blocker: Prophet and recurrent nets need a date index or a 24-hour tensor; /predict has "
        "one feature row. Fix: train them anyway for the comparison table; register only tabular "
        "or ensemble artefacts; log when a non-serveable winner is replaced by the best serveable "
        "model that still beats persistence."
    )
    heading(doc, "15.13 Streamlit and notebook portability", 2)
    body(
        doc,
        "Blockers: emoji-prefixed page filenames were fragile on some systems; EDA notebooks "
        "hard-coded paths; SHAP was a single file instead of per horizon. Fixes: pages named "
        "1_EDA.py … 4_Model_Performance.py; notebooks resolve the repo root from CWD; SHAP written "
        "as shap_summary_{24,48,72}h.png with a legacy fallback on the performance page."
    )

    heading(doc, "16. Limitations and remaining work", 1)
    bullets(
        doc,
        [
            "AQI is PM2.5-only, not the full EPA multi-pollutant maximum.",
            "Approach 1 test months are calm; winter smog performance should be re-checked after the next smog season accumulates under the current OpenWeather regime.",
            "Negative R² at 48 h and 72 h on Approach 1 means those models still beat persistence on error but do not explain test variance; they are “better than no-change,” not “high-skill climate-style R².”",
            "Walk-forward evidence is stronger as a diagnostic than as a registration rule; the gate remains the single chronological test split plus persistence.",
            "Live public hosting of API + Streamlit was not finished at the time of this draft. The local path (uvicorn + streamlit, or docker compose) is implemented. Insert the production URL in the README and in Section 14 when hosting is complete.",
            "Confirm in the Hopsworks UI that aqi_forecaster_24h / 48h / 72h currently hold the Approach 1 artefacts (re-run python -m src.training_pipeline with default TRAIN_START_DATE if a later full-history notebook run left the registry empty).",
            "Commit Approach 1 comparison CSVs and SHAP figures into reports/ so the Streamlit performance page and this report point at the same files.",
        ],
    )

    heading(doc, "17. Conclusion", 1)
    body(
        doc,
        "This project delivers a disciplined FTI forecast of Lahore AQI at one, two and three "
        "days ahead. Pollution and weather each have one source. The target is a continuous EPA "
        "AQI, not a 1–5 vendor category. Features are justified by EDA and computed on a true "
        "hourly grid. Models are allowed into the registry only when they beat persistence on "
        "RMSE, MAE and R²."
    )
    body(
        doc,
        "Approach 1 (data from 4 April 2025) is the selected system: the models beat persistence "
        "and cut AQI-point error roughly in half relative to full-history persistence, with low R² "
        "that is honest for a calm test window. Approach 2 (full history) is documented because "
        "its R² looks stronger, but persistence still wins, so those models are not registered. "
        "A 1–5 classification detour did not overturn that decision. Mentor review agreed with "
        "Approach 1. Remaining close-out work is operational — register the Approach 1 artefacts "
        "if not already current, host the API and dashboard, and attach the live URL — not a "
        "change of modelling strategy."
    )

    heading(doc, "Appendix A. Classification metric tables (1–5 labels)", 1)
    body(
        doc,
        "RMSE, MAE and R² below are computed on OpenWeather classes {1,2,3,4,5}, not on 0–500 AQI. "
        "They are included for completeness of the discarded experiment."
    )
    heading(doc, "A.1 Full history", 2)
    add_table(
        doc,
        ["H", "Model", "RMSE", "MAE", "R²"],
        [
            ["24", "Persistence", "0.729", "0.420", "0.412"],
            ["24", "Ordinal logistic", "1.461", "1.131", "−1.358"],
            ["24", "Random Forest", "1.307", "0.964", "−0.887"],
            ["24", "XGBoost", "1.243", "0.905", "−0.707"],
            ["24", "LightGBM", "1.322", "0.972", "−0.930"],
            ["24", "Ensemble", "1.286", "0.943", "−0.827"],
            ["24", "LSTM", "1.494", "1.155", "−1.462"],
            ["24", "GRU", "1.472", "1.127", "−1.390"],
            ["48", "Persistence", "0.816", "0.500", "0.264"],
            ["48", "Ordinal logistic", "1.512", "1.178", "−1.525"],
            ["48", "Random Forest", "1.493", "1.150", "−1.462"],
            ["48", "XGBoost", "1.353", "1.001", "−1.022"],
            ["48", "LightGBM", "1.334", "0.977", "−0.965"],
            ["48", "Ensemble", "1.392", "1.041", "−1.141"],
            ["48", "LSTM", "1.512", "1.178", "−1.524"],
            ["48", "GRU", "1.503", "1.161", "−1.493"],
            ["72", "Persistence", "0.863", "0.546", "0.178"],
            ["72", "Ordinal logistic", "1.512", "1.176", "−1.525"],
            ["72", "Random Forest", "1.510", "1.173", "−1.519"],
            ["72", "XGBoost", "1.402", "1.049", "−1.171"],
            ["72", "LightGBM", "1.397", "1.050", "−1.157"],
            ["72", "Ensemble", "1.440", "1.092", "−1.291"],
            ["72", "LSTM", "1.521", "1.186", "−1.552"],
            ["72", "GRU", "1.513", "1.173", "−1.525"],
        ],
    )
    caption(doc, "Table A1. Classification, full history. Persistence wins every horizon.")
    heading(doc, "A.2 Regime window", 2)
    add_table(
        doc,
        ["H", "Model", "RMSE", "MAE", "R²"],
        [
            ["24", "Persistence", "0.826", "0.520", "−0.558"],
            ["24", "Ordinal logistic", "0.753", "0.472", "−0.297"],
            ["24", "Random Forest", "0.780", "0.497", "−0.391"],
            ["24", "XGBoost", "0.821", "0.539", "−0.541"],
            ["24", "LightGBM", "0.817", "0.543", "−0.527"],
            ["24", "Ensemble", "0.813", "0.532", "−0.510"],
            ["24", "LSTM", "0.778", "0.495", "−0.379"],
            ["24", "GRU", "0.800", "0.526", "−0.455"],
            ["48", "Persistence", "0.883", "0.602", "−0.789"],
            ["48", "Ordinal logistic", "0.796", "0.524", "−0.456"],
            ["48", "Random Forest", "0.812", "0.534", "−0.513"],
            ["48", "XGBoost", "0.802", "0.540", "−0.476"],
            ["48", "LightGBM", "0.829", "0.560", "−0.578"],
            ["48", "Ensemble", "0.807", "0.539", "−0.496"],
            ["48", "LSTM", "0.797", "0.525", "−0.488"],
            ["48", "GRU", "0.758", "0.480", "−0.346"],
            ["72", "Persistence", "0.895", "0.607", "−0.888"],
            ["72", "Ordinal logistic", "0.776", "0.502", "−0.420"],
            ["72", "Random Forest *", "0.746", "0.474", "−0.311"],
            ["72", "XGBoost", "0.845", "0.570", "−0.683"],
            ["72", "LightGBM", "0.868", "0.594", "−0.777"],
            ["72", "Ensemble", "0.822", "0.548", "−0.592"],
            ["72", "LSTM", "0.799", "0.530", "−0.495"],
            ["72", "GRU", "0.767", "0.489", "−0.380"],
        ],
    )
    caption(doc, "Table A2. Classification, regime window. *Only 72 h Random Forest beat persistence on all five metrics (including accuracy and macro F1).")

    heading(doc, "Appendix B. Repository map", 1)
    bullets(
        doc,
        [
            "src/feature_pipeline.py — raw snapshot, backfill, hourly insert",
            "src/training_pipeline.py — train, gate, register",
            "src/utils/aqi_calculation.py — EPA PM2.5 AQI",
            "src/utils/data_fetch.py — OpenWeather + Open-Meteo",
            "src/utils/feature_engineering.py — hourly grid and features",
            "src/utils/evaluation.py — metrics, shrinkage, persistence gate",
            "api/main.py — FastAPI /predict",
            "app/ — Streamlit",
            "Dockerfile — dashboard and API in one image",
            "notebooks/01_eda.ipynb — EDA on raw parquet",
            "notebooks/02_training.ipynb — Colab regression (Approach 2 run currently saved in outputs)",
            ".github/workflows/ — hourly features, daily training",
            "tests/ — unit tests for the defects in Section 15",
        ],
    )
    body(
        doc,
        "End of report. Edit author name, insert the live demo URL after hosting, and export to PDF "
        "from Microsoft Word (File → Save As → PDF). If Word asks to update fields, that is only "
        "needed if a table of contents field is added later; this draft uses a static Contents list."
    )

    doc.save(OUT)
    print("Wrote", OUT)


if __name__ == "__main__":
    build()
